from flask import Blueprint, request, jsonify, current_app, url_for
from flask_login import login_required, current_user
from app.models import db, User, Notification, Appointment, Message, Game, SessionMetrics, SessionImage, ContactMessage, Sede, Payment
from app.services.appointment_service import AppointmentService
from app.services.game_service import GameService
from app.services.admin_service import AdminService
from app.services.notification_service import NotificationService
from app.services.patient_service import PatientService
from app.services.dashboard_service import DashboardService
from app.services.report_service import ReportService
from app.services.google_drive_service import GoogleDriveService
from app.services.ai_service import predict_level, start_async_training
from app.utils import get_user_today_utc_range, get_user_now, localize_datetime_for_display, get_user_timezone
from app.schemas import AssignTherapistSchema, UpdateUserSchema, SendMessageSchema
from app.extensions import bcrypt, limiter, csrf
from app.services.email_service import EmailService
from app.services.financial_service import FinancialService
from app.utils.api_helpers import api_response
from app.services.availability_service import AvailabilityService
from datetime import datetime, timedelta, timezone
import json
import os
import uuid
from werkzeug.utils import secure_filename
import requests
from sqlalchemy import or_, func

api_bp = Blueprint('api', __name__, url_prefix='/api')

appointment_service = AppointmentService()
game_service = GameService()
admin_service = AdminService()
notification_service = NotificationService()
patient_service = PatientService()
dashboard_service = DashboardService()
report_service = ReportService()
drive_service = GoogleDriveService()
fs = FinancialService()

LIMA_TZ = timezone(timedelta(hours=-5))

def _parse_datetime(value):
    """Parsea fechas como sea. Sin timezone asume Lima (UTC-5). Devuelve UTC pa la BD."""
    if not value:
        return None
    try:
        if value.endswith('Z'):
            value = value[:-1] + '+00:00'
        dt = datetime.fromisoformat(value)
        if dt.tzinfo:
            return dt.astimezone(timezone.utc).replace(tzinfo=None)
        return dt.replace(tzinfo=LIMA_TZ).astimezone(timezone.utc).replace(tzinfo=None)
    except Exception:
        for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%d"):
            try:
                dt = datetime.strptime(value, fmt)
                return dt.replace(tzinfo=LIMA_TZ).astimezone(timezone.utc).replace(tzinfo=None)
            except Exception:
                continue
    return None


@api_bp.route('/time', methods=['GET'])
def api_time():
    now_local = datetime.now()
    now_utc = datetime.utcnow()
    return jsonify({
        'server_time_local': now_local.isoformat(),
        'server_time_utc': now_utc.isoformat(),
        'timezone': 'America/Lima',
        'utc_offset_minutes': -300,
        'is_dst': False,
    })

@api_bp.route('/therapist/insights')
@login_required
def therapist_insights():
    if current_user.role != 'terapista':
        return jsonify({'error': 'Acceso denegado'}), 403
    try:
        data = dashboard_service.get_therapist_insights(current_user)
        return jsonify(data)
    except Exception as e:
        current_app.logger.error(f"Error in therapist_insights: {str(e)}")
        return jsonify({"error": str(e), "data": []}), 500

@api_bp.route('/notifications')
@login_required
def get_notifications():
    notifications = notification_service.get_unread_notifications(current_user.id)
    result = [{
        'id': n.id,
        'title': n.title,
        'type': n.type or 'info',
        'message': n.message,
        'timestamp': n.timestamp.strftime('%d %b, %H:%M'),
        'link': n.link
    } for n in notifications]
    return jsonify(result)

@api_bp.route('/patients')
@login_required
def api_patients():
    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'error': 'Acceso denegado'}), 403
    therapist_id = request.args.get('therapist_id')
    if current_user.role == 'terapista':
        patients = patient_service.get_therapist_patients(current_user.id)
    elif therapist_id and current_user.role == 'admin':
        try:
            from app.models import User
            t_user = User.query.get(int(therapist_id))
            if t_user and t_user.role == 'terapista':
                patients = patient_service.user_repo.get_all_patients_by_therapist(int(therapist_id))
            else:
                patients = []
        except:
            patients = []
    else:
        from app.models import User
        patients = User.query.filter_by(role='jugador', is_active=True).order_by(User.username.asc()).all()
    return jsonify([{'id': p.id, 'username': p.username, 'email': p.email} for p in patients])

@api_bp.route('/notifications/mark-read', methods=['POST'])
@login_required
@csrf.exempt
def mark_notifications_read():
    try:
        data = request.get_json() or {}
        notif_id = data.get('id')
        if notif_id:
            notification_service.mark_one_read(current_user.id, notif_id)
        else:
            notification_service.mark_all_as_read(current_user.id)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@api_bp.route('/sessions', methods=['GET'])
@login_required
def api_get_sessions():
    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'error': 'Acceso denegado'}), 403
    start = request.args.get('start')
    end = request.args.get('end')
    timezone_offset = request.args.get('timezone_offset')
    try:
        start_dt = _parse_datetime(start)
        end_dt = _parse_datetime(end)
    except Exception:
        start_dt = None
        end_dt = None
    if start_dt and end_dt:
        if start_dt.date() == end_dt.date():
            end_dt = end_dt + timedelta(days=1)
        if current_user.role == 'terapista':
            appts = appointment_service.get_therapist_appointments(current_user.id, start_dt, end_dt)
        else:
            appts = appointment_service.get_all_appointments(start_dt, end_dt)
    else:
        q = Appointment.query
        if current_user.role == 'terapista':
            q = q.filter(Appointment.therapist_id == current_user.id)
        appts = q.order_by(Appointment.start_time.desc()).limit(200).all()
    results = []
    for a in appts:
        start_iso = a.start_time.isoformat() if a.start_time else None
        end_iso = a.end_time.isoformat() if a.end_time else None
        results.append({
            'id': a.id,
            'title': a.title or (a.patient.username if a.patient else 'Sesión'),
            'start': start_iso,
            'end': end_iso,
            'status': a.status,
            'attendance': a.attendance,
            'patient': {'id': a.patient.id, 'name': a.patient.username} if a.patient else None,
            'location': a.location,
            'notes': a.notes,
            'games': json.loads(a.games) if a.games else [],
            'is_holiday': True if a.notes and "Scheduled on Holiday" in a.notes else False
        })
    return jsonify(results)

@api_bp.route('/sessions/upcoming', methods=['GET'])
@login_required
def api_upcoming_sessions():
    if current_user.role != 'terapista':
        return jsonify({'error': 'Acceso denegado'}), 403
    appts = appointment_service.get_upcoming_sessions(current_user.id)
    results = []
    for a in appts:
        patient = User.query.get(a.patient_id)
        start_iso = a.start_time.isoformat()
        end_iso = a.end_time.isoformat() if a.end_time else None
        results.append({
            'id': a.id,
            'patient': patient.username or patient.email,
            'start_time': start_iso,
            'end_time': end_iso,
            'status': a.status,
            'attendance': a.attendance,
            'games': json.loads(a.games) if a.games else []
        })
    return jsonify(results)

@api_bp.route('/appointments/patient', methods=['GET'])
@login_required
def api_get_patient_appointments():
    if current_user.role != 'jugador':
        return jsonify({'error': 'Acceso denegado'}), 403
    start = request.args.get('start')
    end = request.args.get('end')
    try:
        start_dt = _parse_datetime(start)
        end_dt = _parse_datetime(end)
    except Exception:
        start_dt = None
        end_dt = None
    appts = appointment_service.get_patient_appointments(current_user.id, start_dt, end_dt)
    results = []
    for a in appts:
        start_iso = a.start_time.isoformat() if a.start_time else None
        end_iso = a.end_time.isoformat() if a.end_time else None
        results.append({
            'id': a.id,
            'title': a.title,
            'start': start_iso,
            'end': end_iso,
            'status': a.status,
            'attendance': a.attendance,
            'therapist': {'id': a.therapist.id, 'name': a.therapist.username} if a.therapist else None,
            'location': a.location,
            'notes': a.notes,
            'games': json.loads(a.games) if a.games else []
        })
    return jsonify(results)


@api_bp.route('/games', methods=['GET'])
@login_required
def api_list_games():
    files = game_service.list_games()
    return jsonify({'games': files})


@api_bp.route('/sessions/day', methods=['GET'])
@login_required
def api_get_sessions_day():

    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'success': False, 'message': 'Acceso denegado'}), 403

    date_str = request.args.get('date')
    timezone_offset = request.args.get('timezone_offset')

    if not date_str:
        return jsonify({'success': False, 'message': 'Falta el parámetro date'}), 400
    try:
        day = _parse_datetime(date_str)
        query_start = day
        query_end = query_start + timedelta(days=1)
    except Exception:
        return jsonify({'success': False, 'message': 'Fecha malita, revisa el formato'}), 400

    # Filter for therapist or admin
    base_query = Appointment.query
    if current_user.role == 'terapista':
        base_query = base_query.filter(Appointment.therapist_id == current_user.id)
    
    query = base_query.filter(Appointment.start_time >= query_start,
                              Appointment.start_time < query_end).order_by(Appointment.start_time.asc()).all()

    results = []
    for a in query:
        start_iso = a.start_time.isoformat()
            
        end_iso = None
        if a.end_time:
            end_iso = a.end_time.isoformat()

        results.append({
            'id': a.id,
            'title': a.title or (a.patient.username if a.patient else 'Sesión'),
            'start': start_iso,
            'end': end_iso,
            'status': a.status,
            'attendance': a.attendance,
            'patient': {'id': a.patient.id, 'name': a.patient.username} if a.patient else None,
            'notes': a.notes,
            'location': a.location
        })

    return jsonify({'date': date_str, 'sessions': results})


@api_bp.route('/sessions', methods=['POST'])
@login_required
def api_create_session():
    """Crear sesión (terapista)"""
    if current_user.role != 'terapista':
        return jsonify({'success': False, 'message': 'Acceso denegado'}), 403

    data = request.json or {}

    for field in ('start_time', 'end_time'):
        if isinstance(data.get(field), str):
            data[field] = _parse_datetime(data[field])

    if not data.get('patient_id') or not data.get('start_time'):
        return jsonify({'success': False, 'message': 'patient_id and start_time are required'}), 400

    if not data.get('end_time'):
        data['end_time'] = data['start_time'] + timedelta(hours=1)
    therapist_id = current_user.id
    if current_user.role == 'admin':
        if 'therapist_ids' in data:
            therapist_id = data['therapist_ids'][0]
        elif 'therapist_id' in data:
            therapist_id = data['therapist_id']
    if therapist_id:
        is_available, error_msg = AvailabilityService.check_availability(
            therapist_id=therapist_id,
            start_time=data['start_time'],
            end_time=data['end_time']
        )
        if not is_available:
            return jsonify({'success': False, 'message': error_msg}), 409
    patient_ids = data.get('patient_id')
    if not isinstance(patient_ids, list):
        patient_ids = [patient_ids]
    if len(patient_ids) > 5:
        return jsonify({'success': False, 'message': 'Máximo 5 pacientes por sesión nomá'}), 400
    created_sessions = []
    all_validation_errors = []
    ignore_therapist_conflict = len(patient_ids) > 1
    for pid in patient_ids:
        validation_errors = appointment_service.validate_session_times(
            start_time=data['start_time'],
            end_time=data['end_time'],
            patient_id=pid,
            therapist_id=current_user.id,
            session_id=None,
            ignore_therapist_conflict=ignore_therapist_conflict
        )
        if validation_errors:
            p_user = User.query.get(pid)
            p_name = p_user.username if p_user else f"ID {pid}"
            for err in validation_errors:
                all_validation_errors.append(f"{p_name}: {err}")
    if all_validation_errors:
        return jsonify({
            'success': False,
            'message': 'Algunos datos no cuadran',
            'errors': all_validation_errors
        }), 400
    try:
        results = []
        for pid in patient_ids:
            session_data = data.copy()
            session_data['patient_id'] = pid
            appt = appointment_service.create_session(current_user.id, session_data, current_user.username)
            created_sessions.append(appt)
            created = {
                'id': appt.id,
                'title': appt.title,
                'start_time': appt.start_time.isoformat() if appt.start_time else None,
                'end_time': appt.end_time.isoformat() if appt.end_time else None,
                'status': appt.status,
                'patient': {'id': appt.patient.id, 'name': appt.patient.username} if appt.patient else None,
                'location': appt.location,
                'notes': appt.notes
            }
            try:
                created['games'] = json.loads(appt.games) if appt.games else []
            except Exception:
                created['games'] = []
            results.append(created)
        return jsonify({
            'success': True,
            'message': 'Sesión registrada, todo ok',
            'session': results[0] if results else {},
            'sessions': results
        }), 201
    except ValueError as e:
        return jsonify({'success': False, 'message': str(e)}), 400
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@api_bp.route('/sessions/<int:session_id>', methods=['GET'])
@login_required
def api_get_session(session_id):
    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'error': 'Acceso denegado'}), 403
    from app.models import SessionImage
    appt = Appointment.query.get_or_404(session_id)
    if current_user.role == 'terapista':
        is_assigned = False
        if appt.patient:
            is_assigned = current_user in appt.patient.therapists
        if appt.therapist_id != current_user.id and not is_assigned:
            return jsonify({'error': 'No tienes permiso para ver esta sesión'}), 403
    images = []
    for img in appt.session_images or []:
        images.append({
            'id': img.id,
            'url': url_for('static', filename=img.image_path),
            'type': img.image_type,
            'notes': img.notes,
            'uploaded_at': img.uploaded_at.isoformat() if img.uploaded_at else None,
            'uploaded_by': img.uploaded_by.username if img.uploaded_by else None,
        })
    start_iso = appt.start_time.isoformat() if appt.start_time else None
    end_iso = appt.end_time.isoformat() if appt.end_time else None
    return jsonify({
        'id': appt.id,
        'title': appt.title or 'Sesión de Terapia',
        'start_time': start_iso,
        'end_time': end_iso,
        'status': appt.status,
        'attendance': appt.attendance,
        'patient': {'id': appt.patient.id, 'name': appt.patient.username} if appt.patient else None,
        'therapist_id': appt.therapist_id,
        'location': appt.location,
        'notes': appt.notes,
        'games': appt.games_list,
        'images': images,
    })

    start_iso = appt.start_time.isoformat() if appt.start_time else None
    end_iso = appt.end_time.isoformat() if appt.end_time else None

    return jsonify({
        'id': appt.id,
        'title': appt.title or 'Sesión de Terapia',
        'start_time': start_iso,
        'end_time': end_iso,
        'status': appt.status,
        'attendance': appt.attendance,
        'patient': {'id': appt.patient.id, 'name': appt.patient.username} if appt.patient else None,
        'therapist_id': appt.therapist_id,
        'location': appt.location,
        'notes': appt.notes,
        'games': appt.games_list,
        'images': images,
    })


@api_bp.route('/sessions/<int:session_id>', methods=['PUT'])
@login_required
def api_update_session(session_id):
    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'success': False, 'message': 'Acceso denegado'}), 403

    data = request.json or {}

    for field in ('start_time', 'end_time'):
        if isinstance(data.get(field), str):
            data[field] = _parse_datetime(data[field])

    existing_appt = Appointment.query.get(session_id)
    if not existing_appt:
        return jsonify({'success': False, 'message': 'Esa sesión no existe'}), 404
    if 'start_time' in data or 'end_time' in data:
        start_time = data.get('start_time', existing_appt.start_time)
        end_time = data.get('end_time', existing_appt.end_time or (existing_appt.start_time + timedelta(hours=1)))
        
        validation_errors = appointment_service.validate_session_times(
            start_time=start_time,
            end_time=end_time,
            patient_id=existing_appt.patient_id,
            therapist_id=current_user.id,
            session_id=session_id
        )
        
        if validation_errors:
            return jsonify({
                'success': False,
                'message': 'Hay observaciones',
                'errors': validation_errors
            }), 400
        
    appt = appointment_service.update_session(session_id, data)
    if not appt:
        return jsonify({'success': False, 'message': 'Esa sesión no existe'}), 404
    updated = {
        'id': appt.id,
        'title': appt.title,
        'start_time': appt.start_time.isoformat() if appt.start_time else None,
        'end_time': appt.end_time.isoformat() if appt.end_time else None,
        'status': appt.status,
        'patient': {'id': appt.patient.id, 'name': appt.patient.username} if appt.patient else None,
        'location': appt.location,
        'attendance': appt.attendance,
        'notes': appt.notes
    }

    return jsonify(updated)


@api_bp.route('/sessions/<int:session_id>', methods=['DELETE'])
@login_required
def api_delete_session(session_id):
    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'success': False, 'message': 'Acceso denegado'}), 403

    success = appointment_service.delete_session(session_id, current_user.id)
    if not success:
        return jsonify({'success': False, 'message': 'Esa sesión no existe'}), 404
    return jsonify({'success': True})

@api_bp.route('/sessions/<int:session_id>/complete', methods=['POST'])
@login_required
def api_complete_session(session_id):

    if current_user.role not in ['terapista', 'admin']:
        return jsonify({'success': False, 'message': 'Acceso denegado'}), 403
    
    try:
        appt = appointment_service.transition_status(
            session_id=session_id,
            new_status='completed',
            changed_by_user_id=current_user.id,
            notify=True
        )
        
        return jsonify({
            'success': True,
            'message': 'Sesión completada, ¡buen trabajo!',
            'session': {
                'id': appt.id,
                'status': appt.status,
                'status_changed_at': appt.status_changed_at.isoformat() if appt.status_changed_at else None
            }
        })
    except ValueError as e:
        return jsonify({'success': False, 'message': str(e)}), 400
    except Exception as e:
        current_app.logger.error(f"Error completing session {session_id}: {str(e)}")
        return jsonify({'success': False, 'message': 'No se pudo completar la sesión'}), 500

@api_bp.route('/sessions/<int:session_id>/cancel', methods=['POST'])
@login_required
def api_cancel_session(session_id):

    if current_user.role not in ['terapista', 'admin']:
        return jsonify({'success': False, 'message': 'Acceso denegado'}), 403
    
    data = request.get_json(silent=True) or {}
    reason = data.get('reason', '')
    
    try:
        appt = appointment_service.transition_status(
            session_id=session_id,
            new_status='cancelled',
            changed_by_user_id=current_user.id,
            notify=True
        )
        
        # Optionally store cancellation reason in notes
        if reason:
            if appt.notes:
                appt.notes += f"\n\n[Cancelada] {reason}"
            else:
                appt.notes = f"[Cancelada] {reason}"
            db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Sesión cancelada, listo',
            'session': {
                'id': appt.id,
                'status': appt.status,
                'status_changed_at': appt.status_changed_at.isoformat() if appt.status_changed_at else None
            }
        })
    except ValueError as e:
        return jsonify({'success': False, 'message': str(e)}), 400
    except Exception as e:
        current_app.logger.error(f"Error cancelling session {session_id}: {str(e)}")
        return jsonify({'success': False, 'message': 'No se pudo cancelar la sesión'}), 500

@api_bp.route('/admin/assign-therapist', methods=['POST'])
@login_required
def api_admin_assign_therapist():
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': 'Acceso denegado'}), 403
    data = request.get_json(silent=True) or {}
    errors = AssignTherapistSchema().validate(data)
    if errors:
        return jsonify({'success': False, 'message': 'Datos inválidos', 'errors': errors}), 400
    
    # Support multiple therapists
    success, message = False, "Error desconocido"
    
    if 'therapist_ids' in data:
        success, message = admin_service.assign_therapist(data['patient_id'], therapist_ids=data['therapist_ids'])
    else:
        # Legacy fallback
        success, message = admin_service.assign_therapist(data['patient_id'], therapist_id=data.get('therapist_id'))

    if not success:
        return jsonify({'success': False, 'message': message}), 400
        
    return jsonify({'success': True})

@api_bp.route('/admin/create-user', methods=['POST'])
@login_required
def api_admin_create_user():
    try:
        if current_user.role != 'admin':
            return jsonify({'success': False, 'message': 'Acceso denegado'}), 403
        data = request.get_json(silent=True) or {}
        
        # Validation logic updated to allow "Patient without email" (Presencial)
        role = data.get('role')
        email = data.get('email', '').strip()
        username = data.get('username', '').strip()

        if not role:
            return jsonify({'success': False, 'message': 'El rol es obligatorio'}), 400
        
        # If role is NOT patient, email is mandatory
        if role != 'jugador' and not email:
            return jsonify({'success': False, 'message': 'El email es obligatorio para administradores y terapeutas'}), 400
            
        # If role IS patient, either email OR username is required
        if role == 'jugador' and not email and not username:
            return jsonify({'success': False, 'message': 'Debes ingresar al menos el Nombre del paciente'}), 400

        success, result = admin_service.create_user(data)
        if not success:
            return jsonify({'success': False, 'message': result}), 400
            
        user_obj = result.get('user') if isinstance(result, dict) else None
        temp_pass = result.get('temp_password') if isinstance(result, dict) else None
        
        if not user_obj:
            # Fallback if service returned weird format
            return jsonify({'success': True, 'message': 'Usuario creado (sin datos de retorno)'})
            
        return jsonify({
            'success': True, 
            'message': 'Usuario creado',
            'temp_password': temp_pass
        })
    except Exception as e:
        current_app.logger.error(f"Error creating user: {str(e)}")
        return jsonify({'success': False, 'message': f"Server Error: {str(e)}"}), 500

@api_bp.route('/admin/reset-password', methods=['POST'])
@login_required
def api_admin_reset_password():
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': 'Acceso denegado'}), 403
    data = request.get_json(silent=True) or {}
    user_id = data.get('id')
    new_password = data.get('new_password')
    
    if not user_id:
        return jsonify({'success': False, 'message': 'ID de usuario requerido'}), 400
    
    success, result = admin_service.reset_user_password(user_id, new_password)
    if not success:
        return jsonify({'success': False, 'message': result}), 400
        
    return jsonify({'success': True, 'temp_password': result})

@api_bp.route('/admin/games/delete', methods=['POST'])
@login_required
def api_admin_delete_game():
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': 'Acceso denegado'}), 403
    data = request.get_json(silent=True) or {}
    name = (data.get('name') or '').strip()
    if not name:
        return jsonify({'success': False, 'message': 'El nombre es obligatorio'}), 400
    games_dir = os.path.join(current_app.root_path, 'static', 'games')
    path = os.path.join(games_dir, name)
    try:
        if os.path.isfile(path):
            os.remove(path)
            return jsonify({'success': True})
        else:
            return jsonify({'success': False, 'message': 'Ese archivo no está'}), 404
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@api_bp.route('/admin/messages/broadcast', methods=['POST'])
@login_required
def api_admin_broadcast():
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': 'Acceso denegado'}), 403
    data = request.get_json(silent=True) or {}
    subject = (data.get('subject') or '').strip()
    body = (data.get('body') or '').strip()
    target = (data.get('target') or 'all').strip()
    receiver_id = data.get('receiver_id')
    
    if not body:
        return jsonify({'success': False, 'message': 'El mensaje no puede estar vacío'}), 400
        
    success, result = admin_service.broadcast_message(current_user.id, subject, body, target, receiver_id)
    if not success:
        return jsonify({'success': False, 'message': result}), 404
        
    return jsonify({'success': True, 'count': result})

def _serialize_user(u):
    return {
        'id': u.id,
        'email': u.email,
        'username': u.username,
        'role': u.role,
        'is_active': u.is_active,
        'account_status': u.account_status or 'active',
        'admin_password_changed_count': u.admin_password_changed_count or 0,
        'sede_id': u.sede_id,
        'sede_name': u.sede_item.name if u.sede_item else None,
        'assigned_sedes': [{'id': s.id, 'name': s.name} for s in u.assigned_sedes.all()],
        'therapist_ids': [t.id for t in u.therapists.all()],
        'payment_plan': u.payment_plan,
        'payment_amount': u.payment_amount or 0,
        'sessions_total': u.sessions_total or 0,
        'sessions_attended': u.sessions_attended or 0,
        'plan_type': u.plan_type or 'individual',
        'has_second_shift': u.has_second_shift or False,
        'payment_amount_2': u.payment_amount_2 or 0,
        'sessions_total_2': u.sessions_total_2 or 0,
        'sessions_attended_2': u.sessions_attended_2 or 0,
        'plan_type_2': u.plan_type_2 or 'individual',
        'salary_base': u.salary_base or 0,
        'contract_hours': u.contract_hours or 0,
        'work_start_time': u.work_start_time,
        'work_end_time': u.work_end_time,
        'work_days': u.work_days,
    }

@api_bp.route('/admin/list-users')
@login_required
def api_admin_list_users():
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': 'Acceso denegado'}), 403
    role = (request.args.get('role') or '').strip()
    users = admin_service.list_users(role)
    return jsonify({'success': True, 'users': [_serialize_user(u) for u in users]})

@api_bp.route('/admin/user/<int:user_id>')
@login_required
def api_admin_get_user(user_id):
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': 'Acceso denegado'}), 403
    u = User.query.get(user_id)
    if not u:
        return jsonify({'success': False, 'message': 'Ese usuario no existe'}), 404
    return jsonify({'success': True, 'user': _serialize_user(u)})

@api_bp.route('/admin/update-user', methods=['POST'])
@login_required
def api_admin_update_user():
    try:
        if current_user.role != 'admin':
            return jsonify({'success': False, 'message': 'Acceso denegado'}), 403
            
        data = request.get_json(silent=True) or {}
        
        errors = UpdateUserSchema().validate(data)
        if errors:
            return jsonify({'success': False, 'message': 'Datos inválidos', 'errors': errors}), 400
            
        success, result = admin_service.update_user(data)
        if not success:
            return jsonify({'success': False, 'message': result}), 400
            
        return jsonify({'success': True})
    except Exception as e:
        current_app.logger.error(f"Error updating user: {str(e)}")
        return jsonify({'success': False, 'message': f"Server Error: {str(e)}"}), 500

@api_bp.route('/admin/delete-user', methods=['POST'])
@login_required
def api_admin_delete_user():
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': 'Acceso denegado'}), 403
    data = request.get_json(silent=True) or {}
    user_id = data.get('id')
    if not user_id:
        return jsonify({'success': False, 'message': 'ID requerido'}), 400
    u = User.query.get(user_id)
    if not u:
        return jsonify({'success': False, 'message': 'Ese usuario no existe'}), 404
    if u.email == (os.getenv('ADMIN_EMAIL') or 'diegocenteno537@gmail.com'):
        return jsonify({'success': False, 'message': 'No puedes borrar al admin principal'}), 400
    try:
        # Cascade delete dependencies first (Explicit for safety)
        Message.query.filter((Message.sender_id==u.id)|(Message.receiver_id==u.id)).delete()
        Appointment.query.filter((Appointment.therapist_id==u.id)|(Appointment.patient_id==u.id)).delete()
        SessionMetrics.query.filter(SessionMetrics.user_id==u.id).delete()
        Payment.query.filter(Payment.patient_id==u.id).delete()
        
        db.session.delete(u)
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500

@api_bp.route('/save_game', methods=['POST'])
@login_required
@limiter.limit("20 per minute")
def save_game():
    try:
        data = request.get_json() or {}
        game_name = data.get('game_name') or 'Juego'
        accuracy = float(data.get('accuracy') or 0)
        avg_time = float(data.get('avg_time') or 0)
        session_id = data.get('session_id')
        
        appt = None
        if session_id:
            appt = Appointment.query.get(int(session_id))
            if not appt:
                return jsonify({'error': 'Esa sesión no existe'}), 404
            
            if current_user.role == 'jugador' and appt.patient_id != current_user.id:
                return jsonify({'error': 'No autorizado para esta sesión'}), 403
            
            if appt.status == 'completed':
                return jsonify({'error': 'Esta sesión ya se completó'}), 400
                
            # 3. Game Assignment check (Optional but recommended)
            # We check if the game being saved is actually part of the session
            # Using the new property games_list that handles both legacy and new models
            # We normalize names for comparison (remove .html, case insensitive)
            assigned_normalized = [g.lower().replace('.html', '').replace('_', ' ') for g in appt.games_list]
            current_normalized = game_name.lower().replace('.html', '').replace('_', ' ')
            
            # Permitir si la lista esta vacia (legacy/testing) o si hay match
            if appt.games_list and current_normalized not in assigned_normalized:
                current_app.logger.warning(f"Game mismatch: {game_name} not in {appt.games_list}")
                # return jsonify({'error': 'Juego no asignado a esta sesión'}), 400

        pred_code, label = predict_level(accuracy, avg_time * 1000)  # avg_time expected in seconds; convert ms for model input

        m = SessionMetrics(
            user_id=current_user.id,
            session_id=int(session_id) if session_id else None,
            game_name=game_name,
            accurracy=accuracy,
            avg_time=avg_time,
            prediction=pred_code
        )
        
        # Link to Game model if possible
        # Try to find game by filename or title
        game_obj = Game.query.filter(or_(Game.filename == game_name, Game.title == game_name)).first()
        if game_obj:
            m.game_id = game_obj.id
            
        db.session.add(m)

        # If tied to a session, check progress
        if appt:
            # Flush to ensure the new metric is countable
            db.session.flush()
            
            # Count total metrics for this session
            played_count = SessionMetrics.query.filter_by(session_id=appt.id).count()
            
            # Only close if we have played at least as many games as assigned
            # Use the robust games_list property
            total_assigned = len(appt.games_list)
            
            if played_count >= total_assigned and total_assigned > 0:
                appt.status = 'completed'
                appt.end_time = datetime.utcnow()
                db.session.add(appt)
                # Optional: create a notification for therapist
                try:
                    notification_service.create_notification(appt.therapist_id, f"Sesión #{appt.id} completada por {current_user.username}", link=url_for('therapist.patients', _external=False))
                except Exception:
                    pass


        db.session.commit()

        try:
            total_metrics = SessionMetrics.query.count()
            if total_metrics > 0 and total_metrics % 5 == 0:
                all_metrics = SessionMetrics.query.all()
                training_data = [[m.accurracy, m.avg_time * 1000] for m in all_metrics]
                current_app.logger.info(f"Triggering AI async retraining with {len(training_data)} samples...")
                start_async_training(training_data)
        except Exception as e:
            current_app.logger.error(f"AI Retraining trigger failed: {e}")

        return jsonify({'status': 'ok', 'prediction': pred_code, 'recommendation': label})
    except Exception as e:
        return jsonify({'error': 'no_se_pudo_guardar', 'detail': str(e)}), 400


@api_bp.route('/games/upload', methods=['POST'])
@limiter.limit("5 per hour")
@login_required
def upload_game():
    if current_user.role != 'terapista':
        return jsonify({'error': 'Acceso denegado'}), 403
    file = request.files.get('file')
    name = request.form.get('name')
    if not file or not name:
        return jsonify({'error': 'Falta el archivo o el nombre'}), 400
    if not name.lower().endswith('.html'):
        name = f"{name}.html"
    dest_dir = os.path.join(current_app.root_path, 'static', 'games')
    os.makedirs(dest_dir, exist_ok=True)
    path = os.path.join(dest_dir, name)
    file.save(path)
    return jsonify({'status': 'ok', 'file': name, 'url': url_for('static', filename=f'games/{name}')})


@api_bp.route('/ai/gemini', methods=['POST'])
@limiter.limit("10 per minute")
@login_required
def gemini_proxy():
    if current_user.role not in ('terapista','admin'):
        return jsonify({'error': 'Acceso denegado'}), 403
    api_key = current_app.config.get('GEMINI_API_KEY')
    payload = request.get_json() or {}
    prompt = payload.get('prompt')
    context = payload.get('context')
    if not prompt:
        return jsonify({'error': 'Falta el prompt'}), 400
    if not api_key:
        acc = (context or {}).get('accuracy') or 0
        avg = (context or {}).get('avg_time') or 0
        _, label = predict_level(acc, avg)
        return jsonify({'status': 'no_external', 'recommendation': label})
    try:
        import requests
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-pro:generateContent?key={api_key}"
        headers = {'Content-Type': 'application/json'}
        data = {
            "contents": [{
                "parts": [{"text": f"Context: {json.dumps(context)}. Prompt: {prompt}"}]
            }]
        }
        resp = requests.post(url, headers=headers, json=data)
        if resp.status_code == 200:
            result = resp.json()
            text = result.get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text', '')
            return jsonify({'status': 'ok', 'response': text})
        else:
            return jsonify({'error': 'Gemini falló', 'details': resp.text}), 500
    except Exception as e:
        return jsonify({'error': 'Gemini no respondió', 'detail': str(e)}), 500

@api_bp.route('/ai/generate_game', methods=['POST'])
@login_required
def generate_game():
    if current_user.role not in ('terapista','admin'):
        return jsonify({'error': 'Acceso denegado'}), 403
    api_key = current_app.config.get('GEMINI_API_KEY')
    payload = request.get_json() or {}
    prompt = payload.get('prompt') or 'Genera un juego terapéutico en HTML.'
    target_user_id = payload.get('user_id')
    game_name = (payload.get('name') or 'ai_game').strip().replace(' ', '_')
    if not target_user_id:
        return jsonify({'error': 'Falta user_id'}), 400
    user = User.query.get(target_user_id)
    if not user:
        return jsonify({'error': 'Usuario no encontrado'}), 404

    kpi = {}
    kpi['total_sessions'] = SessionMetrics.query.filter_by(user_id=user.id).count()
    kpi['avg_accuracy'] = float(db.session.query(func.avg(SessionMetrics.accurracy)).filter_by(user_id=user.id).scalar() or 0)
    kpi['avg_time_ms'] = float((db.session.query(func.avg(SessionMetrics.avg_time)).filter_by(user_id=user.id).scalar() or 0) * 1000)
    kpi['last_games'] = [
        {
            'game_name': m.game_name,
            'accuracy': float(m.accurracy),
            'avg_time_ms': float(m.avg_time * 1000),
            'prediction': int(m.prediction),
            'date': m.date.isoformat()
        } for m in SessionMetrics.query.filter_by(user_id=user.id).order_by(SessionMetrics.date.desc()).limit(10)
    ]

    full_prompt = (
        f"{prompt}\n\n"
        "Genera dos bloques: 1) HTML completo para un juego sencillo de reflejos/cognitivo con UI moderna, tailwindcdn y FontAwesome (no frameworks).\n"
        "2) JSON de configuración KPI con claves: kpis(avg_accuracy, avg_time_ms, total_sessions), goals, difficulty, and tracking schema for events.\n"
        f"KPIs del paciente: {json.dumps(kpi, ensure_ascii=False)}\n"
        "Devuelve primero el JSON (entre marcadores ---JSON---) y luego el HTML (entre ---HTML---)."
    )

    if not api_key:
        # Fallback: simple generated HTML and JSON locally
        config = {
            'kpis': {'avg_accuracy': kpi['avg_accuracy'], 'avg_time_ms': kpi['avg_time_ms'], 'total_sessions': kpi['total_sessions']},
            'goals': ['Mejorar reflejos', 'Reducir tiempo de reacción'],
            'difficulty': 'medium',
            'tracking': {'events': ['click', 'hit', 'miss'], 'schema_version': 1}
        }
        html = '<!DOCTYPE html><html><head><meta charset="utf-8"><script src="https://cdn.tailwindcss.com"></script></head><body class="p-6">\n' \
               '<h2 class="text-2xl font-bold">Juego IA (fallback)</h2>\n' \
               '<p class="text-gray-600">Config basado en KPIs.</p>\n' \
               '</body></html>'
    else:
        try:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-pro:generateContent?key={api_key}"
            body = {"contents": [{"parts": [{"text": full_prompt}]}]}
            resp = requests.post(url, json=body, timeout=15)
            resp.raise_for_status()
            j = resp.json()
            text = (
                j.get('candidates', [{}])[0]
                .get('content', {})
                .get('parts', [{}])[0]
                .get('text') or ''
            )
            json_start = text.find('---JSON---')
            html_start = text.find('---HTML---')
            if json_start != -1 and html_start != -1:
                json_block = text[json_start + len('---JSON---'): html_start].strip()
                html_block = text[html_start + len('---HTML---'):].strip()
                try:
                    config = json.loads(json_block)
                except Exception:
                    config = {'raw': json_block}
                html = html_block
            else:
                # If markers missing, store raw
                config = {'raw': text}
                html = '<!DOCTYPE html><html><body><pre>Salida IA sin marcadores</pre></body></html>'
        except Exception as e:
            config = {'error': str(e), 'kpis': kpi}
            html = '<!DOCTYPE html><html><body><pre>Error generando juego IA</pre></body></html>'

    dest_dir = os.path.join(current_app.root_path, 'static', 'games')
    os.makedirs(dest_dir, exist_ok=True)
    filename = f"{game_name}.html"
    path = os.path.join(dest_dir, filename)
    try:
        with open(path, 'w', encoding='utf-8') as f:
            f.write(html)
    except Exception as e:
        return jsonify({'error': 'write_failed', 'detail': str(e)}), 500

    try:
        user.game_profile = json.dumps(config, ensure_ascii=False)
        db.session.commit()
    except Exception as e:
        return jsonify({'error': 'persist_failed', 'detail': str(e)}), 500

    return jsonify({
        'status': 'ok',
        'file': filename,
        'url': url_for('static', filename=f'games/{filename}'),
        'config': config
    })

@api_bp.route('/sessions/assign-games', methods=['POST'])
@login_required
def assign_games_to_session():
    """Asignar juegos vía AppointmentGame"""
    if current_user.role not in ('terapista','admin'):
        return jsonify({'error': 'Acceso denegado'}), 403
    
    data = request.get_json() or {}
    session_id = data.get('session_id')
    games = data.get('games') or []
    
    if not session_id:
        return jsonify({'error': 'session_id requerido'}), 400
    
    # Extract filenames from games list
    # Support both ['game.html'] and [{'name': 'game.html', 'url': '...'}]
    game_filenames = []
    for game in games:
        if isinstance(game, dict):
            game_filenames.append(game.get('name', ''))
        else:
            game_filenames.append(game)
    
    # Filter out empty strings
    game_filenames = [g for g in game_filenames if g]
    
    try:
        validated_games = appointment_service.set_session_games(session_id, game_filenames)
        return jsonify({
            'status': 'ok',
            'assigned': [{'name': g.filename, 'title': g.title} for g in validated_games]
        })
    except ValueError as e:
        return jsonify({'error': str(e)}), 404
    except Exception as e:
        return jsonify({'error': f'Error asignando juegos: {str(e)}'}), 500


# Check available games for a session (only during time window)
@api_bp.route('/sessions/<int:session_id>/games', methods=['GET'])
@login_required
def session_games(session_id):
    appt = Appointment.query.get(session_id)
    if not appt:
        return jsonify({'error': 'Sesión no encontrada'}), 404
    now = datetime.utcnow()
    # Allow access if now is between start and end (or within scheduled with end None -> 2h)
    end_time = appt.end_time or (appt.start_time + timedelta(hours=2))
    enabled = appt.status == 'scheduled' and appt.start_time <= now <= end_time
    games = []
    try:
        games = json.loads(appt.games) if appt.games else []
    except Exception:
        games = []
    return jsonify({'enabled': enabled, 'games': games})


# Aggregate session results and update user profile (game_profile)
@api_bp.route('/sessions/<int:session_id>/complete', methods=['POST'])
@login_required
def complete_session(session_id):
    appt = Appointment.query.get(session_id)
    if not appt:
        return jsonify({'error': 'Sesión no encontrada'}), 404
    if current_user.id != appt.therapist_id:
        return jsonify({'error': 'Acceso denegado'}), 403

    if appt.status != 'completed':
        appt.status = 'completed'
        appt.end_time = datetime.utcnow()
        db.session.add(appt)

    metrics = SessionMetrics.query.filter_by(user_id=appt.patient_id, session_id=session_id).all()
    if not metrics:
        db.session.commit()
        return jsonify({'status': 'ok', 'message': 'Sin métricas para agregar'})

    avg_acc = float(sum(m.accurracy for m in metrics) / len(metrics))
    avg_time_ms = float(sum(m.avg_time for m in metrics) / len(metrics) * 1000)
    plays = len(metrics)
    last_games = [{
        'game_name': m.game_name,
        'accuracy': float(m.accurracy),
        'avg_time_ms': float(m.avg_time * 1000),
        'prediction': int(m.prediction),
        'date': m.date.isoformat()
    } for m in metrics]

    # Merge into user.game_profile JSON
    patient = User.query.get(appt.patient_id)
    try:
        existing = json.loads(patient.game_profile) if patient.game_profile else {}
    except Exception:
        existing = {}
    existing.setdefault('history', []).extend(last_games)
    existing['kpis'] = {
        'avg_accuracy': avg_acc,
        'avg_time_ms': avg_time_ms,
        'plays': plays
    }

    patient.game_profile = json.dumps(existing, ensure_ascii=False)
    db.session.commit()

    try:
        notification_service.create_notification(appt.therapist_id, f"Sesión #{appt.id} completada. {plays} juegos registrados.", link=url_for('therapist.reports'))
        notification_service.create_notification(appt.patient_id, f"Sesión completada. ¡Buen trabajo!", link=url_for('patient.progress'))
    except Exception:
        pass

    return jsonify({'status': 'ok', 'updated_profile': existing})

@api_bp.route('/resources/<int:resource_id>')
@login_required
def get_resource(resource_id):
    try:
        if resource_id == 1:
            # Guía de Ejercicios: summarize recent performance
            metrics = SessionMetrics.query.filter_by(user_id=current_user.id).order_by(SessionMetrics.date.desc()).limit(20).all()
            if metrics:
                avg_acc = sum((m.accurracy or 0) for m in metrics) / len(metrics)
                avg_time = sum((m.avg_time or 0) for m in metrics) / len(metrics)
                perf_summary = f"Tu precisión promedio en las últimas sesiones es {avg_acc:.0f}%. Tiempo medio por ejercicio {avg_time:.1f}s."
            else:
                perf_summary = "No hay datos de sesiones suficientes para personalizar esta guía."

            content = f"<h3>Guía de Ejercicios Personalizada</h3><p>{perf_summary}</p>"
            content += "<ol><li>Ejercicio respiratorio: 5 minutos.</li><li>Ejercicios de atención: 3 bloques de 4 minutos.</li><li>Revisión de estrategias aprendidas en la sesión.</li></ol>"
            return jsonify({'id': resource_id, 'title': 'Guía de Ejercicios', 'content': content})

        if resource_id == 2:
            content = "<h3>Video Tutorial: Técnicas básicas</h3><p>Este video explica las técnicas recomendadas y cuándo aplicarlas. Duración: 15:30.</p>"
            content += "<p>Puntos clave: respiración, pausas activas, seguimiento de progreso.</p>"
            return jsonify({'id': resource_id, 'title': 'Video Tutorial', 'content': content})

        if resource_id == 3:
            content = "<h3>Hoja de Práctica</h3><p>Plantilla descargable para llevar un registro de ejercicios diarios.</p>"
            content += "<ul><li>Día 1: Ejercicio A - 10 repeticiones</li><li>Día 2: Ejercicio B - 8 repeticiones</li></ul>"
            return jsonify({'id': resource_id, 'title': 'Hoja de Práctica', 'content': content})

        return jsonify({'error': 'Recurso no encontrado'}), 404
    except Exception as e:
        return jsonify({'error': 'Error generando recurso', 'detail': str(e)}), 500

@api_bp.route('/messages/send', methods=['POST'])
@login_required
@csrf.exempt
def send_message():
    if request.is_json:
        data = request.get_json()
    else:
        data = request.form.to_dict()

    receiver_id = data.get('receiver_id')
    body = data.get('body', '')
    subject = data.get('subject')
    
    if not receiver_id:
        return jsonify({'success': False, 'message': 'Datos incompletos'}), 400
    
    receiver = User.query.get(receiver_id)
    if not receiver:
        return jsonify({'success': False, 'message': 'Destinatario no encontrado'}), 404
    
    attachment_path = None
    attachment_type = None
    
    if 'file' in request.files:
        file = request.files['file']
        if file and file.filename:
            filename = secure_filename(file.filename)
            unique_filename = f"{uuid.uuid4().hex}_{filename}"
            
            ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
            if ext in ['jpg', 'jpeg', 'png', 'gif', 'webp']:
                attachment_type = 'image'
            elif ext in ['mp4', 'mov', 'webm']:
                attachment_type = 'video'
            elif ext in ['mp3', 'wav', 'ogg', 'm4a']:
                attachment_type = 'audio'
            else:
                attachment_type = 'file'

            upload_folder = os.path.join(current_app.instance_path, 'uploads', 'messages')
            os.makedirs(upload_folder, exist_ok=True)
            
            file.save(os.path.join(upload_folder, unique_filename))
            
            # Storing filename only, path is managed by frontend/template
            attachment_path = unique_filename

    if not body and not attachment_path:
         return jsonify({'success': False, 'message': 'El mensaje no puede estar vacío'}), 400
    
    message = Message(
        sender_id=current_user.id,
        receiver_id=receiver_id,
        subject=subject,
        body=body,
        attachment_path=attachment_path,
        attachment_type=attachment_type
    )
    db.session.add(message)
    db.session.commit()
    
    notification_service.create_notification(
        user_id=receiver_id,
        title=f'Nuevo mensaje de {current_user.username}',
        message=body or 'Has recibido un archivo adjunto',
        notif_type='message',
        link=url_for('main.messages_list')
    )
    
    try:
        email_service = EmailService()
        email_service.send_new_message_email(
            receiver.email,
            receiver.username,
            current_user.username,
            (body or "Has recibido un archivo adjunto")[:100] + ('...' if body and len(body) > 100 else '')
        )
    except Exception:
        pass # Non-critical
    
    return jsonify({
        'success': True,
        'message_id': message.id,
        'created_at': message.created_at.isoformat(),
        'attachment_path': attachment_path,
        'attachment_type': attachment_type
    })


@api_bp.route('/messages/unread-count')
@login_required
def unread_messages_count():
    count = Message.query.filter_by(receiver_id=current_user.id, is_read=False).count()
    return jsonify({'count': count})

@api_bp.route('/admin/profile', methods=['POST'])
@login_required
def api_admin_update_profile():
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': 'Acceso denegado'}), 403
    data = request.get_json(silent=True) or {}
    name = (data.get('username') or '').strip()
    new_password = (data.get('new_password') or '').strip()
    changed = False
    if name:
        current_user.username = name
        changed = True
    if new_password:
        current_user.password = bcrypt.generate_password_hash(new_password).decode('utf-8')
        changed = True
        try:
            EmailService.send_password_change_email(current_user.email, new_password, current_user.username or 'Administrador')
        except Exception:
            pass
    if changed:
        db.session.commit()
    return jsonify({'success': True})

@api_bp.route('/appointments/<int:appointment_id>/upload_image', methods=['POST'])
@login_required
def upload_session_image(appointment_id):

    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'error': 'Acceso denegado'}), 403
        
    appointment = Appointment.query.get_or_404(appointment_id)
    
    if current_user.role == 'terapista' and appointment.therapist_id != current_user.id:
        return jsonify({'error': 'No tienes permiso para editar esta sesión'}), 403
        
    if 'image' not in request.files:
        return jsonify({'error': 'No se encontró el archivo de imagen'}), 400
        
    file = request.files['image']
    image_type = request.form.get('image_type', 'session_photo')
    notes = request.form.get('notes', '')
    
    if file.filename == '':
        return jsonify({'error': 'Nombre de archivo vacío'}), 400
        
    if file:
        allowed_extensions = {'png', 'jpg', 'jpeg', 'webp', 'doc', 'docx'}
        if '.' not in file.filename or \
           file.filename.rsplit('.', 1)[1].lower() not in allowed_extensions:
            return jsonify({'error': 'Tipo de archivo no permitido (solo imágenes y Word)'}), 400
            
        # Create secure filename with UUID to prevent collisions
        original_filename = secure_filename(file.filename)
        extension = original_filename.rsplit('.', 1)[1].lower()
        unique_filename = f"{uuid.uuid4().hex}.{extension}"
        
        # Create directory structure: static/uploads/session_images/YYYY/MM
        now = datetime.utcnow()
        relative_path = os.path.join('uploads', 'session_images', str(now.year), f"{now.month:02d}")
        upload_folder = os.path.join(current_app.root_path, 'static', relative_path)
        
        os.makedirs(upload_folder, exist_ok=True)
        
        file_path = os.path.join(upload_folder, unique_filename)
        file.save(file_path)

        # Upload to Google Drive (Background/Sync)
        try:
            # Pass the file PATH instead of opening it, to avoid stream issues
            patient_name = appointment.patient.username if appointment.patient else 'Paciente_Desconocido'
            session_date = appointment.start_time.strftime('%Y-%m-%d')
            
            print(f"Subiendo a Drive: {patient_name} / {session_date} / {unique_filename}")
            drive_service.upload_file(
                file_path,  # Path string 
                unique_filename,
                file.mimetype,
                patient_name,
                session_date
            )
        except Exception as e:
            print(f"Error subiendo a Google Drive: {str(e)}")
        
        # Store relative path in DB for serving
        db_relative_path = os.path.join(relative_path, unique_filename)
        
        session_image = SessionImage(
            appointment_id=appointment.id,
            image_path=db_relative_path,
            image_type=image_type,
            uploaded_by_id=current_user.id,
            notes=notes
        )
        
        db.session.add(session_image)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'image': {
                'id': session_image.id,
                'url': url_for('static', filename=db_relative_path),
                'type': session_image.image_type,
                'notes': session_image.notes
            }
        })
        
    return jsonify({'error': 'Error al subir archivo'}), 500

@api_bp.route('/appointments/<int:appointment_id>/images/<int:image_id>', methods=['DELETE'])
@login_required
def delete_session_image(appointment_id, image_id):

    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'error': 'Acceso denegado'}), 403

    image = SessionImage.query.get_or_404(image_id)
    
    # Verify it belongs to the appointment
    if image.appointment_id != appointment_id:
        return jsonify({'error': 'Imagen no corresponde a la sesión'}), 400
        
    # Verify ownership (therapist assigned to appointment)
    if current_user.role == 'terapista':
        appointment = Appointment.query.get(appointment_id)
        if appointment.therapist_id != current_user.id:
            return jsonify({'error': 'No tienes permiso para editar esta sesión'}), 403

    try:
        full_path = os.path.join(current_app.root_path, 'static', image.image_path)
        if os.path.exists(full_path):
            os.remove(full_path)
            
        db.session.delete(image)
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': f'Error al eliminar imagen: {str(e)}'}), 500

# --- SEDES MANAGEMENT ---
from app.models import Sede

@api_bp.route('/admin/sedes', methods=['GET', 'POST'])
@login_required
def admin_sedes():
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': 'Forbidden'}), 403

    try:
        if request.method == 'GET':
            sedes = Sede.query.order_by(Sede.created_at.desc()).all()
            result = []
            for s in sedes:
                created_at_iso = None
                if s.created_at:
                    try:
                        created_at_iso = s.created_at.isoformat()
                    except AttributeError:
                        created_at_iso = str(s.created_at)
                
                result.append({
                    'id': s.id,
                    'name': s.name,
                    'address': s.address,
                    'active': s.active,
                    'created_at': created_at_iso
                })
            return jsonify(result)

        if request.method == 'POST':
            data = request.get_json() or {}
            name = data.get('name')
            if not name:
                return jsonify({'success': False, 'message': 'Nombre es obligatorio'}), 400
            
            existing = Sede.query.filter_by(name=name).first()
            if existing:
                return jsonify({'success': False, 'message': 'Sede ya existe'}), 400

            address = data.get('address')
            s = Sede(name=name, address=address)
            db.session.add(s)
            db.session.commit()
            return jsonify({'success': True, 'id': s.id})
    except Exception as e:
        current_app.logger.error(f"Error in admin_sedes: {str(e)}")
        return jsonify({"error": str(e), "data": []}), 500

@api_bp.route('/admin/sedes/<int:sede_id>', methods=['PUT', 'GET'])
@login_required
def admin_sedes_detail(sede_id):
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
        
    try:
        s = Sede.query.get(sede_id)
        if not s:
            return jsonify({'success': False, 'message': 'No encontrado'}), 404

        if request.method == 'PUT':
            data = request.get_json() or {}
            if 'active' in data:
                s.active = bool(data['active'])
            if 'name' in data and data['name']:
                s.name = data['name']
            if 'address' in data:
                s.address = data['address']
            
            db.session.commit()
            return jsonify({'success': True})
        
        return jsonify({
            'id': s.id,
            'name': s.name,
            'address': s.address,
            'active': s.active
        })
    except Exception as e:
        current_app.logger.error(f"Error in admin_sedes_detail: {str(e)}")
        return jsonify({"error": str(e), "data": []}), 500

@api_bp.route('/admin/sedes/<int:sede_id>/analytics', methods=['GET'])
@login_required
def admin_sedes_analytics(sede_id):
    """Analíticas de sede"""
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    
    try:
        sede = Sede.query.get(sede_id)
        if not sede:
            return jsonify({'success': False, 'message': 'Sede not found'}), 404
        
        from datetime import datetime
        # Fix division by zero or empty list issues in statistics
        # Get therapists assigned to this sede
        therapists = User.query.filter(
            User.assigned_sedes.any(Sede.id == sede_id),
            User.role == 'terapista'
        ).all()
        therapist_ids = [t.id for t in therapists]
        
        # Get all patients (jugador role) who have appointments with these therapists
        appointments_at_sede = Appointment.query.filter(
            Appointment.therapist_id.in_(therapist_ids)
        ).all() if therapist_ids else []
        
        patient_ids = list(set([a.patient_id for a in appointments_at_sede if a.patient_id]))
        
        # Payments for patients at this sede
        payments = Payment.query.filter(
            Payment.patient_id.in_(patient_ids)
        ).all() if patient_ids else []
        
        total_patients = len(patient_ids)
        active_patients = len([pid for pid in patient_ids if User.query.get(pid) and User.query.get(pid).is_active])
        
        total_revenue = sum([p.amount for p in payments if p.status == 'completed']) if payments else 0
        total_sessions = len([a for a in appointments_at_sede if a.status == 'completed'])
        pending_sessions = len([a for a in appointments_at_sede if a.status == 'scheduled'])
        
        today = datetime.utcnow()
        month_start = datetime(today.year, today.month, 1)
        sessions_this_month = len([a for a in appointments_at_sede if a.status == 'completed' and a.start_time and a.start_time >= month_start])
        payments_this_month = sum([p.amount for p in payments if p.status == 'completed' and p.date and p.date >= month_start]) if payments else 0
        
        # Therapists assigned to this sede (already fetched as therapists)
        
        return jsonify({
            'success': True,
            'sede': {
                'id': sede.id,
                'name': sede.name,
                'address': sede.address,
            },
            'analytics': {
                'patients': {
                    'total': total_patients,
                    'active': active_patients,
                },
                'payments': {
                    'total_revenue': round(total_revenue, 2),
                    'this_month': round(payments_this_month, 2),
                    'transactions': len(payments),
                },
                'sessions': {
                    'total_completed': total_sessions,
                    'pending': pending_sessions,
                    'this_month': sessions_this_month,
                    'total': len(appointments_at_sede),
                },
                'therapists': {
                    'count': len(therapists),
                    'names': [t.email for t in therapists],
                }
            }
        })
    except Exception as e:
        current_app.logger.error(f"Error in admin_sedes_analytics: {str(e)}")
        return jsonify({"error": str(e), "data": []}), 500

@api_bp.route('/admin/deudores', methods=['GET'])
@login_required
def admin_deudores_por_sede():
    """Reporte de deuda delegado a FinancialService"""
    if current_user.role != 'admin':
        return jsonify({"error": "Forbidden", "data": []}), 403

    month = request.args.get('month', 'all')
    if month == 'curr':
        month = 'current'
    try:
        data = fs.build_debt_report(days_ahead=7, month=month)
        # Ensure 'por_sede' exists even if empty
        if not data or 'por_sede' not in data:
            data = {"por_sede": {}, "summary": {}}
        return api_response(success=True, data=data)
    except Exception as e:
        current_app.logger.error(f"Financial report failed: {str(e)}")
        import traceback
        current_app.logger.error(traceback.format_exc())
        return api_response(success=False, error=str(e), data={"por_sede": {}}, status=500)

@api_bp.route('/v1/payments/<int:payment_id>/mark-paid', methods=['POST'])
@login_required
def mark_payment_paid(payment_id):
    """Marcar pago como completado y reactivar"""
    if current_user.role != 'admin':
        return jsonify({'error': 'Unauthorized', 'message': 'Solo admins pueden realizar esta acción.'}), 403
        
    payment = Payment.query.get_or_404(payment_id)
    
    data = request.get_json() or {}
    method = data.get('method', payment.method or 'transfer')
    
    payment.status = 'completed'
    payment.method = method
    payment.date = datetime.utcnow()
    
    # También activar al usuario si estaba inactivo por falta de pago
    if payment.patient_id:
        user = User.query.get(payment.patient_id)
        if user and not user.is_active:
            user.is_active = True
            
    db.session.commit()
    
    return jsonify({
        'success': True,
        'message': f'Pago de {payment.amount} registrado exitosamente.',
        'payment_id': payment.id
    })

@api_bp.route('/admin/send-payment-reminder', methods=['POST'])
@login_required
def send_payment_reminder():
    """Enviar recordatorio de pago"""
    if current_user.role != 'admin':
        return jsonify({'success': False, 'error': 'Unauthorized'}), 403
    
    try:
        data = request.get_json() or {}
        patient_id = data.get('patient_id')
        patient_email = data.get('patient_email')
        channel = data.get('channel', 'email')
        
        if not patient_id or not patient_email:
            return jsonify({'success': False, 'error': 'patient_id and patient_email required'}), 400
        
        from app.services.financial_service import FinancialService
        fs = FinancialService()
        info = fs.get_patient_overdue_info(patient_id)
        if not info:
            return api_response(success=False, error={'message': 'Patient not found'}, status=404)

        due_date = info.get('due_date')
        amount = info.get('amount', 0)
        days_overdue = info.get('days_overdue', 0)
        patient_name = info.get('name')
        
        if channel == 'email':
            from app.services.email_service import EmailService
            
            subject = f"Recordatorio: pago pendiente - Centro de Terapias"
            body = f"""
Hola {patient_name},

Te escribimos para recordarte que tienes una deuda pendiente.

Detalles:
- Monto: S/ {amount:.2f}
- Vencimiento: {due_date.strftime('%d/%m/%Y') if due_date else 'N/A'}
- Días de atraso: {days_overdue}

Por favor, ponte al día para evitar acciones adicionales.

Si ya pagaste, haz como que no viste este mensaje.

Gracias,
Centro de Terapias
"""
            try:
                EmailService.send_email(patient_email, subject, body)
                return api_response(success=True, data={'message': f'Recordatorio enviado a {patient_email}', 'channel': 'email'})
            except Exception as e:
                current_app.logger.error(f"Error sending email reminder: {e}")
                return api_response(success=False, error={'message': f'Error al enviar recordatorio por email: {str(e)}'}, status=500)
        
        elif channel in ['sms', 'whatsapp']:
            # Try SMS/WhatsApp via Twilio
            from app.services.sms_whatsapp_service import SMSWhatsAppService
            
            sms_service = SMSWhatsAppService()
            
            if not sms_service.is_available():
                return jsonify({
                    'success': False,
                    'error': 'Servicio SMS/WhatsApp no disponible. Configure Twilio credentials.'
                }), 501
            
            # Use data.get('phone') or fetch it from User model if missing
            from app.models import User
            patient_record = User.query.get(patient_id)
            phone_number = (patient_record.phone if patient_record else None) or data.get('phone')
            if not phone_number:
                return jsonify({
                    'success': False,
                    'error': 'No phone number available for this patient'
                }), 400
            
            # Send via SMS or WhatsApp
            if channel == 'sms':
                success = sms_service.send_payment_reminder_sms(
                    phone_number, patient_name, amount, due_date, days_overdue
                )
            else:  # whatsapp
                success = sms_service.send_payment_reminder_whatsapp(
                    phone_number, patient_name, amount, due_date, days_overdue
                )
            
                if success:
                    return api_response(success=True, data={'message': f'Recordatorio enviado por {channel} a {phone_number}', 'channel': channel})
                else:
                    return api_response(success=False, error={'message': f'Error al enviar recordatorio por {channel}'}, status=500)
        
        else:
            return api_response(success=False, error={'message': f'Canal no soportado: {channel}. Use email, sms, o whatsapp.'}, status=400)
        
    except Exception as e:
        current_app.logger.error(f"Unexpected error in send_payment_reminder: {e}")
        import traceback
        traceback.print_exc()
        return api_response(success=False, error={'message': str(e)}, status=500)

@api_bp.route('/v1/search/patients', methods=['GET'])
@login_required
def search_patients():
    """Búsqueda global de pacientes para el Command+K Modal"""
    if current_user.role not in ['admin', 'therapist']:
        return jsonify({'error': 'Unauthorized'}), 403
        
    query = request.args.get('q', '').strip()
    if len(query) < 2:
        return jsonify({'patients': []})
        
    # Search by username, email, phone...
    search_term = f"%{query}%"
    patients = User.query.filter(
        db.or_(
            User.username.ilike(search_term),
            User.email.ilike(search_term),
            User.phone.ilike(search_term) if hasattr(User, 'phone') else db.false()
        ),
        User.role == 'jugador'
    ).limit(10).all()
    
    result = []
    for p in patients:
        result.append({
            'id': p.id,
            'username': p.username,
            'email': p.email,
            'phone': getattr(p, 'phone', '')
        })
        
    return jsonify({'patients': result})

@api_bp.route('/public/contact', methods=['POST'])
@csrf.exempt
def contact_message():
    data = request.get_json() or {}
    required_fields = ['first_name', 'last_name', 'email', 'message']
    for field in required_fields:
        if field not in data or not data[field]:
            return jsonify({'error': f'Falta el campo {field}'}), 400

    new_msg = ContactMessage(
        first_name=data.get('first_name'),
        last_name=data.get('last_name'),
        email=data.get('email'),
        phone=data.get('phone'),
        subject=data.get('subject', 'Consulta Web'),
        message=data.get('message'),
        service_interest=data.get('service_interest'),
        urgency=data.get('urgency', 'medium'),
        status='unread'
    )
    
    try:
        db.session.add(new_msg)
        db.session.commit()
        return jsonify({'message': '¡Mensaje recibido! Te contactamos pronto.'}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@api_bp.route('/admin/metrics/capacity', methods=['GET'])
@login_required
def get_capacity_metrics():

    if current_user.role != 'admin':
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        capacity_data = admin_metrics_service.get_capacity_metrics()
        therapist_load = admin_metrics_service.get_therapist_load()
        user_health = admin_metrics_service.get_user_health_kpi()
        
        return jsonify({
            'success': True,
            'capacity': capacity_data,
            'therapist_load': therapist_load,
            'user_health': user_health
        })
    except Exception as e:
        current_app.logger.error(f"Error fetching capacity metrics: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@api_bp.route('/notifications/create', methods=['POST'])
@login_required
@csrf.exempt
def create_notification():
    try:
        data = request.get_json()
        message = data.get('message')
        title = data.get('title')
        notif_type = data.get('type', 'info')
        link = data.get('link')

        if not message:
            return jsonify({'success': False, 'message': 'Mensaje es requerido'}), 400

        notification_service.create_notification(
            user_id=current_user.id,
            title=title,
            message=message,
            notif_type=notif_type,
            link=link
        )
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


# ——— AUDITORÍA IA — Whisper + Llama 3 ———

@api_bp.route('/sessions/<int:appointment_id>/program', methods=['POST'])
@login_required
def upload_session_program(appointment_id):
    """Subir .docx de programación de sesión"""
    if current_user.role != 'admin':
        return jsonify({'success': False, 'error': 'Solo el administrador puede subir la programación'}), 403

    appointment = Appointment.query.get_or_404(appointment_id)

    if 'program_file' not in request.files:
        return jsonify({'success': False, 'error': 'No se encontró el archivo'}), 400

    file = request.files['program_file']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'Nombre de archivo vacío'}), 400

    if not file.filename.lower().endswith('.docx'):
        return jsonify({'success': False, 'error': 'Solo se aceptan archivos .docx'}), 400

    try:
        from app.services.audit_service import extract_docx_text
        from app.models import SessionAudit

        temp_filename = f"temp_program_{uuid.uuid4().hex}.docx"
        temp_dir = os.path.join(current_app.config.get('UPLOAD_FOLDER', '/tmp'), 'temp_audit')
        os.makedirs(temp_dir, exist_ok=True)
        temp_path = os.path.join(temp_dir, temp_filename)
        file.save(temp_path)

        try:
            planned_text = extract_docx_text(temp_path)
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

        audit = SessionAudit.query.filter_by(appointment_id=appointment_id).first()
        if not audit:
            audit = SessionAudit(appointment_id=appointment_id)
            db.session.add(audit)

        audit.planned_text = planned_text
        audit.docx_uploaded_at = datetime.utcnow()
        audit.docx_uploaded_by = current_user.id
        # Resetear auditoría si se sube nueva programación
        audit.audit_status = 'pending'
        audit.audit_report_json = None
        audit.audit_score = None

        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Programación subida correctamente',
            'planned_text_preview': planned_text[:500] + ('...' if len(planned_text) > 500 else ''),
            'char_count': len(planned_text)
        })

    except ValueError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        current_app.logger.error(f"Error subiendo programación: {str(e)}")
        return jsonify({'success': False, 'error': f'Error interno: {str(e)}'}), 500


@api_bp.route('/sessions/auto-complete-expired', methods=['POST'])
@login_required
def completar_sesiones_vencidas():
    cutoff = datetime.now(timezone.utc) - timedelta(minutes=5)
    expired = Appointment.query.filter(
        Appointment.status == 'in_progress',
        Appointment.end_time < cutoff
    ).all()
    count = 0
    for appt in expired:
        appt.status = 'completed'
        count += 1
    db.session.commit()
    return jsonify({'success': True, 'completed': count})


@api_bp.route('/sessions/<int:appointment_id>/audio', methods=['POST'])
@login_required
@csrf.exempt
def upload_session_audio(appointment_id):
    """Subir audio para transcripción Whisper (se elimina tras transcribir)"""
    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

    appointment = Appointment.query.get_or_404(appointment_id)

    if current_user.role == 'terapista' and appointment.therapist_id != current_user.id:
        return jsonify({'success': False, 'error': 'No tienes permiso para esta sesión'}), 403

    if 'audio_file' not in request.files:
        return jsonify({'success': False, 'error': 'No se encontró el archivo de audio'}), 400

    file = request.files['audio_file']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'Nombre de archivo vacío'}), 400

    allowed_audio = {'webm', 'wav', 'mp3', 'ogg', 'm4a', 'mp4'}
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in allowed_audio:
        return jsonify({'success': False, 'error': f'Formato no soportado. Usa: {", ".join(allowed_audio)}'}), 400

    try:
        from app.services.audit_service import transcribe_audio
        from app.models import SessionAudit

        # Guardar temporalmente (será eliminado por audit_service)
        temp_filename = f"session_audio_{appointment_id}_{uuid.uuid4().hex}.{ext}"
        temp_dir = os.path.join(current_app.config.get('UPLOAD_FOLDER', '/tmp'), 'temp_audio')
        os.makedirs(temp_dir, exist_ok=True)
        temp_path = os.path.join(temp_dir, temp_filename)
        file.save(temp_path)

        # Transcribir (el archivo se elimina dentro de transcribe_audio)
        result = transcribe_audio(temp_path)

        audit = SessionAudit.query.filter_by(appointment_id=appointment_id).first()
        if not audit:
            audit = SessionAudit(appointment_id=appointment_id)
            db.session.add(audit)

        # Append transcript (supports chunked recording every 5 min)
        existing = audit.transcript_text or ''
        separator = ' ' if existing else ''
        audit.transcript_text = existing + separator + result['text']
        audit.audio_transcribed_at = datetime.utcnow()
        audit.audio_duration_seconds = (audit.audio_duration_seconds or 0) + result.get('duration', 0)
        if audit.audit_status == 'completed':
            audit.audit_status = 'pending'
            audit.audit_report_json = None
            audit.audit_score = None

        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Audio transcrito correctamente. Archivo eliminado del servidor.',
            'transcript_text': result['text'],
            'transcript_preview': result['text'][:500] + ('...' if len(result['text']) > 500 else ''),
            'duration_seconds': result.get('duration', 0),
            'char_count': len(result['text'])
        })

    except ValueError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        current_app.logger.error(f"Error transcribiendo audio: {str(e)}")
        # Asegurar eliminación del audio en caso de error no manejado
        try:
            if 'temp_path' in locals() and os.path.exists(temp_path):
                os.remove(temp_path)
                current_app.logger.info(f" Audio eliminado tras error: {temp_path}")
        except Exception:
            pass
        return jsonify({'success': False, 'error': f'Error al transcribir: {str(e)}'}), 500


@api_bp.route('/sessions/<int:appointment_id>/audit', methods=['POST'])
@login_required
def trigger_session_audit(appointment_id):
    """Disparar auditoría IA: programación vs transcripción"""
    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

    try:
        from app.services.audit_service import run_audit
        from app.models import User
        report = run_audit(appointment_id)
        appt = Appointment.query.get(appointment_id)
        score = None
        if report:
            score = report.get('score') or (report.get('report') or {}).get('audit_score') if isinstance(report, dict) else None
        therapist_name = appt.therapist.username if appt and appt.therapist else 'Desconocido'
        patient_name = appt.patient.username if appt and appt.patient else 'Desconocido'
        score_str = f' — Puntuación: {score}/100' if score is not None else ''
        from app.services.notification_service import NotificationService
        ns = NotificationService()
        msg_admin = f'Auditoría completada: {therapist_name} / {patient_name}{score_str}'
        admins = User.query.filter_by(role='admin').all()
        for admin in admins:
            ns.create_notification(admin.id, msg_admin)
        msg_therapist = f'Auditoría completada para {patient_name}{score_str}'
        if appt and appt.therapist_id:
            ns.create_notification(appt.therapist_id, msg_therapist)

        # Auto-generar reporte diario después de la auditoría
        if appt and score is not None:
            try:
                from app.services.report_service import ReportService
                rs = ReportService()
                session_date = appt.start_time.date() if appt.start_time else datetime.utcnow().date()
                rs.generate_daily_report(appt.patient_id, appt.therapist_id, session_date.isoformat())
            except Exception as daily_err:
                current_app.logger.error(f"Error generando reporte diario post-auditoría: {daily_err}")

        return jsonify({
            'success': True,
            'message': 'Auditoría completada',
            'report': report
        })
    except ValueError as e:
        err_msg = str(e)
        current_app.logger.warning(f"Audit ValueError for session {appointment_id}: {err_msg}")
        return jsonify({'success': False, 'error': err_msg, 'reason': 'validation'}), 400
    except Exception as e:
        current_app.logger.error(f"Error en auditoría IA: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': f'Error en auditoría: {str(e)}', 'reason': 'server'}), 500


@api_bp.route('/sessions/<int:appointment_id>/audit', methods=['GET'])
@login_required
def get_session_audit(appointment_id):
    """Estado y reporte de auditoría de sesión"""
    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

    from app.models import SessionAudit

    audit = SessionAudit.query.filter_by(appointment_id=appointment_id).first()
    if not audit:
        return jsonify({
            'success': True,
            'exists': False,
            'message': 'No hay registro de auditoría para esta sesión'
        })

    return jsonify({
        'success': True,
        'exists': True,
        'audit': {
            'id': audit.id,
            'has_program': bool(audit.planned_text),
            'has_transcript': bool(audit.transcript_text),
            'planned_text_preview': (audit.planned_text[:300] + '...') if audit.planned_text and len(audit.planned_text) > 300 else audit.planned_text,
            'transcript_preview': (audit.transcript_text[:300] + '...') if audit.transcript_text and len(audit.transcript_text) > 300 else audit.transcript_text,
            'planned_text': audit.planned_text,
            'transcript_text': audit.transcript_text,
            'audio_duration_seconds': audit.audio_duration_seconds,
            'audit_status': audit.audit_status,
            'audit_score': audit.audit_score,
            'report': audit.get_report() if audit.audit_status == 'completed' else None,
            'docx_uploaded_at': audit.docx_uploaded_at.isoformat() if audit.docx_uploaded_at else None,
            'audio_transcribed_at': audit.audio_transcribed_at.isoformat() if audit.audio_transcribed_at else None,
            'audited_at': audit.audited_at.isoformat() if audit.audited_at else None
        }
    })


@api_bp.route('/sessions/<int:appointment_id>/compare-live', methods=['GET'])
@login_required
def compare_session_live(appointment_id):
    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

    from app.models import SessionAudit
    from app.services.audit_service import compute_similarity_vectorial

    audit = SessionAudit.query.filter_by(appointment_id=appointment_id).first()
    if not audit:
        return jsonify({'success': False, 'error': 'No hay auditoría'}), 404

    vectorial = compute_similarity_vectorial(audit.planned_text or '', audit.transcript_text or '')

    duracion = audit.audio_duration_seconds or 0
    ratio = min(1.0, duracion / 2700)
    factor = min(1.0, ratio / 0.1)

    return jsonify({
        'success': True,
        'score_vectorial': vectorial['score_vectorial'],
        'objetivos_cubiertos': vectorial['objetivos_cubiertos'],
        'n_objectives': vectorial['n_objectives'],
        'ratio_duracion': round(ratio, 3),
        'factor_penalizacion': round(factor, 3),
        'duracion_segundos': duracion,
        'char_count': len(audit.transcript_text or '')
    })


@api_bp.route('/sessions/<int:appointment_id>/report-docx', methods=['GET'])
@login_required
def download_report_docx(appointment_id):
    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

    from app.models import SessionAudit, Appointment
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    appt = Appointment.query.get(appointment_id)
    if not appt:
        return jsonify({'success': False, 'error': 'Sesión no encontrada'}), 404

    if current_user.role == 'terapista' and appt.therapist_id != current_user.id:
        return jsonify({'success': False, 'error': 'No tienes acceso a esta sesión'}), 403

    audit = SessionAudit.query.filter_by(appointment_id=appointment_id).first()
    if not audit or audit.audit_status != 'completed':
        return jsonify({'success': False, 'error': 'No hay auditoría completada para esta sesión'}), 400

    report = audit.get_report()
    patient_name = appt.patient.username if appt.patient else '—'
    therapist_name = appt.therapist.username if appt.therapist else '—'

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    title = doc.add_heading('Reporte de Auditoría IA', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Sesión: {appt.title or "Sin título"}')
    doc.add_paragraph(f'Paciente: {patient_name}')
    doc.add_paragraph(f'Terapeuta: {therapist_name}')
    doc.add_paragraph(f'Fecha: {appt.start_time.strftime("%d/%m/%Y %H:%M") if appt.start_time else "—"}')
    doc.add_paragraph(f'Puntaje: {audit.audit_score or "—"} / 100')
    doc.add_paragraph(f'Clasificación: {report.get("status", "—")}')

    doc.add_heading('Objetivos Terapéuticos', level=1)
    objectives = report.get('objectives', [])
    if objectives:
        for obj in objectives:
            p = doc.add_paragraph()
            run = p.add_run(f'{obj.get("name", "—")}')
            run.bold = True
            p.add_run(f'\n  Evidencia: {obj.get("evidence", "—")}')
            p.add_run(f'\n  Clasificación: {obj.get("classification", "—")}')
    else:
        doc.add_paragraph('No se registraron objetivos.')

    if report.get('observations'):
        doc.add_heading('Observaciones IA', level=1)
        doc.add_paragraph(report['observations'])

    if audit.transcript_text:
        doc.add_heading('Transcripción', level=1)
        doc.add_paragraph(audit.transcript_text[:2000] + ('...' if len(audit.transcript_text) > 2000 else ''))

    feedback_items = []
    if audit.feedback_engagement:
        feedback_items.append(f'Compromiso: {audit.feedback_engagement}/5')
    if audit.feedback_progress:
        feedback_items.append(f'Progreso: {audit.feedback_progress}/5')
    if audit.feedback_notes:
        feedback_items.append(f'Notas: {audit.feedback_notes}')
    if feedback_items:
        doc.add_heading('Feedback del Terapeuta', level=1)
        for item in feedback_items:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph('Generado por EdySync', style='Intense Quote')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    patient_safe = patient_name.replace(' ', '_').lower()
    filename = f'auditoria_{patient_safe}_{appointment_id}.docx'

    from flask import send_file
    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )


@api_bp.route('/sessions/<int:appointment_id>/program', methods=['DELETE'])
@login_required
def delete_session_program(appointment_id):
    """Eliminar programación .docx"""
    if current_user.role != 'admin':
        return jsonify({'success': False, 'error': 'Solo el administrador puede eliminar la programación'}), 403

    from app.models import SessionAudit
    audit = SessionAudit.query.filter_by(appointment_id=appointment_id).first()
    if not audit or not audit.planned_text:
        return jsonify({'success': False, 'error': 'No hay programación para esta sesión'}), 404

    audit.planned_text = None
    audit.docx_uploaded_at = None
    audit.docx_uploaded_by = None
    # Reset audit if it was completed
    if audit.audit_status == 'completed':
        audit.audit_status = 'pending'
        audit.audit_report_json = None
        audit.audit_score = None
        audit.audited_at = None

    # If no transcript either, delete the whole record
    if not audit.transcript_text:
        db.session.delete(audit)
    
    db.session.commit()
    return jsonify({'success': True, 'message': 'Programación eliminada'})


@api_bp.route('/sessions/<int:appointment_id>/program', methods=['GET'])
@login_required
def get_session_program(appointment_id):
    """Texto de programación para terapista"""
    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

    from app.models import SessionAudit
    audit = SessionAudit.query.filter_by(appointment_id=appointment_id).first()
    if not audit or not audit.planned_text:
        return jsonify({'success': False, 'exists': False})

    return jsonify({
        'success': True,
        'exists': True,
        'planned_text': audit.planned_text,
        'uploaded_at': audit.docx_uploaded_at.isoformat() if audit.docx_uploaded_at else None
    })




# ——— REPORTES — Semanal x paciente + Daily ———

@api_bp.route('/reports/generate-weekly', methods=['POST'])
@login_required
def generate_weekly_report():

    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

    data = request.get_json(silent=True) or {}
    patient_id = data.get('patient_id')
    week_start = data.get('week_start')

    if not patient_id or not week_start:
        return jsonify({'success': False, 'error': 'patient_id y week_start son requeridos'}), 400

    try:
        from app.services.report_service import ReportService
        rs = ReportService()
        therapist_id = current_user.id if current_user.role == 'terapista' else data.get('therapist_id', current_user.id)
        report = rs.generate_patient_weekly_report(patient_id, therapist_id, week_start)
        return jsonify({'success': True, 'report': report})
    except ValueError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        current_app.logger.error(f"Error generating weekly report: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@api_bp.route('/reports/weekly/<int:patient_id>', methods=['GET'])
@login_required
def get_weekly_report(patient_id):
    if current_user.role not in ('terapista', 'admin', 'jugador'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

    week_start = request.args.get('week')

    try:
        from app.services.report_service import ReportService
        rs = ReportService()
        report = rs.get_patient_weekly_report(patient_id, week_start)
        if not report:
            return jsonify({'success': True, 'exists': False})
        return jsonify({'success': True, 'exists': True, 'report': report})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@api_bp.route('/reports/generate-daily', methods=['POST'])
@login_required
def generate_daily_report():

    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

    data = request.get_json(silent=True) or {}
    patient_id = data.get('patient_id')
    report_date = data.get('date')

    if not patient_id:
        return jsonify({'success': False, 'error': 'patient_id es requerido'}), 400

    try:
        from app.services.report_service import ReportService
        rs = ReportService()
        therapist_id = current_user.id if current_user.role == 'terapista' else data.get('therapist_id', current_user.id)
        report = rs.generate_daily_report(patient_id, therapist_id, report_date)
        return jsonify({'success': True, 'report': report})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

# ——— NO-SHOW / ASISTENCIA ———

@api_bp.route('/sessions/<int:session_id>/start-recording', methods=['POST'])
@login_required
def start_session_recording(session_id):
    """Marcar sesión como en_progreso para grabación"""
    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

    appt = Appointment.query.get_or_404(session_id)
    if current_user.role == 'terapista' and appt.therapist_id != current_user.id:
        return jsonify({'success': False, 'error': 'No tienes permiso para esta sesión'}), 403

    if appt.status not in ('scheduled', 'in_progress', 'completed'):
        return jsonify({'success': False, 'error': f'Estado de sesión inválido: {appt.status}'}), 400

    if appt.status != 'in_progress':
        appt.status = 'in_progress'
        appt.status_changed_at = datetime.utcnow()
        appt.status_changed_by = current_user.id
        db.session.commit()

    return jsonify({'success': True, 'message': 'Grabación iniciada', 'session_id': session_id})


@api_bp.route('/sessions/<int:session_id>/analyze-attendance', methods=['POST'])
@login_required
def analyze_session_attendance(session_id):
    """Detectar inasistencia vía transcripción vs plan"""
    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

    appt = Appointment.query.get_or_404(session_id)
    if current_user.role == 'terapista' and appt.therapist_id != current_user.id:
        return jsonify({'success': False, 'error': 'No tienes permiso para esta sesión'}), 403

    from app.models import SessionAudit
    audit = SessionAudit.query.filter_by(appointment_id=session_id).first()

    transcript = audit.transcript_text if audit and audit.transcript_text else ''
    planned = audit.planned_text if audit and audit.planned_text else ''

    if not transcript or len(transcript.strip()) < 50:
        return jsonify({
            'success': True,
            'suggested_attendance': 'absent',
            'confidence': 0.95,
            'reason': 'Sin transcripción o muy corta',
            'coverage_pct': 0
        })

    if not planned:
        return jsonify({
            'success': True,
            'suggested_attendance': 'present',
            'confidence': 0.5,
            'reason': 'Sin programación para comparar',
            'coverage_pct': 50
        })

    try:
        from app.services.audit_service import analyze_attendance
        result = analyze_attendance(planned, transcript)
        return jsonify({
            'success': True,
            'suggested_attendance': result['suggested_attendance'],
            'confidence': result['confidence'],
            'reason': result.get('reason', ''),
            'coverage_pct': result.get('coverage_pct', 0)
        })
    except ValueError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        current_app.logger.error(f"Error analyzing attendance: {str(e)}")
        return jsonify({'success': False, 'error': f'Error: {str(e)}'}), 500


@api_bp.route('/sessions/<int:session_id>/mark-absent', methods=['POST'])
@login_required
def mark_session_absent(session_id):
    """Marcar sesión como ausente"""
    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

    appt = Appointment.query.get_or_404(session_id)
    if current_user.role == 'terapista' and appt.therapist_id != current_user.id:
        return jsonify({'success': False, 'error': 'No tienes permiso para esta sesión'}), 403

    appt.attendance = 'absent'
    appt.status = 'completed'
    appt.status_changed_at = datetime.utcnow()
    appt.status_changed_by = current_user.id
    db.session.commit()

    return jsonify({
        'success': True,
        'message': 'Sesión marcada como ausente',
        'session_id': session_id
    })


# ═══════════════════════════════════════════════════════════════════
# THERAPIST FEEDBACK
# ═══════════════════════════════════════════════════════════════════

@api_bp.route('/sessions/<int:session_id>/feedback', methods=['POST'])
@login_required
def submit_session_feedback(session_id):
    """Feedback del terapeuta sobre la sesión"""
    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

    appt = Appointment.query.get_or_404(session_id)
    if current_user.role == 'terapista' and appt.therapist_id != current_user.id:
        return jsonify({'success': False, 'error': 'No tienes permiso para esta sesión'}), 403

    data = request.get_json(silent=True) or {}
    engagement = data.get('engagement')
    progress = data.get('progress')
    notes = data.get('notes', '')

    from app.models import SessionAudit
    audit = SessionAudit.query.filter_by(appointment_id=session_id).first()
    if not audit:
        audit = SessionAudit(appointment_id=session_id)
        db.session.add(audit)

    audit.feedback_engagement = engagement
    audit.feedback_progress = progress
    audit.feedback_notes = notes
    audit.feedback_submitted_at = datetime.utcnow()
    db.session.commit()

    return jsonify({'success': True, 'message': 'Feedback guardado'})


@api_bp.route('/therapist/efficiency', methods=['GET'])
@login_required
def get_therapist_efficiency():
    """Eficiencia del terapeuta"""
    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

    therapist_id = current_user.id
    if current_user.role == 'admin':
        therapist_id = request.args.get('therapist_id', type=int) or current_user.id

    try:
        from app.services.dashboard_service import DashboardService
        ds = DashboardService()
        efficiency = ds.get_therapist_efficiency(therapist_id)
        return jsonify({'success': True, 'data': efficiency})
    except Exception as e:
        current_app.logger.error(f"Error getting efficiency: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@api_bp.route('/admin/audit-stats', methods=['GET'])
@login_required
def get_audit_stats():
    if current_user.role != 'admin':
        return jsonify({'error': 'Unauthorized'}), 403
        
    try:
        from app.models import SessionAudit, Appointment, User
        from sqlalchemy import func
        from app.extensions import db

        total_audits = SessionAudit.query.filter(SessionAudit.audit_score.isnot(None)).count()
        avg_score_q = db.session.query(func.avg(SessionAudit.audit_score)).filter(SessionAudit.audit_score.isnot(None)).scalar()
        avg_score = float(avg_score_q) if avg_score_q else 0.0
        
        recent_audits = db.session.query(
            SessionAudit.id,
            SessionAudit.audit_score,
            SessionAudit.audited_at,
            Appointment.title,
            User.username.label('therapist_name')
        ).join(
            Appointment, SessionAudit.appointment_id == Appointment.id
        ).join(
            User, Appointment.therapist_id == User.id
        ).filter(
            SessionAudit.audit_score.isnot(None)
        ).order_by(SessionAudit.audited_at.desc()).limit(10).all()
        
        audit_rows = []
        for r in recent_audits:
            audit_rows.append({
                'id': r[0],
                'score': float(r[1]),
                'date': r[2].isoformat() if r[2] else None,
                'title': r[3],
                'therapist': r[4]
            })

        therapist_scores = db.session.query(
            User.username,
            func.avg(SessionAudit.audit_score),
            func.count(SessionAudit.id)
        ).join(
            Appointment, SessionAudit.appointment_id == Appointment.id
        ).join(
            User, Appointment.therapist_id == User.id
        ).filter(
            SessionAudit.audit_score.isnot(None)
        ).group_by(User.id).all()
        
        therapist_stats = []
        for t in therapist_scores:
            therapist_stats.append({
                'name': t[0],
                'avg_score': round(float(t[1]), 1),
                'count': t[2]
            })

        return jsonify({
            'success': True,
            'data': {
                'total': total_audits,
                'avg_score': round(avg_score, 1),
                'recent': audit_rows,
                'by_therapist': therapist_stats
            }
        })
    except Exception as e:
        current_app.logger.warning(f"Failed to load audit stats: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@api_bp.route('/therapist/weekly-reports/pending', methods=['GET'])
@login_required
def api_weekly_reports_pending():
    if current_user.role != 'terapista':
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403
    try:
        from app.models import WeeklyReport, Notification
        week_start = datetime.utcnow().date()
        monday = week_start - timedelta(days=week_start.weekday())
        reports = WeeklyReport.query.filter(
            WeeklyReport.therapist_id == current_user.id,
            WeeklyReport.week_start == monday
        ).count()
        notification = Notification.query.filter(
            Notification.user_id == current_user.id,
            Notification.type == 'reportes',
            Notification.is_read == False
        ).order_by(Notification.created_at.desc()).first()
        return jsonify({
            'success': True,
            'has_pending': reports > 0,
            'reports_count': reports,
            'has_notification': notification is not None,
            'week_start': monday.isoformat(),
            'week_end': (monday + timedelta(days=6)).isoformat(),
        })
    except Exception as e:
        current_app.logger.error(f"Error in weekly-reports/pending: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': 'Error al consultar reportes'}), 500

@api_bp.route('/weekly-summary', methods=['GET'])
@login_required
def api_weekly_summary():
    if current_user.role != 'admin':
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403
    try:
        week_start = request.args.get('week_start')
        if not week_start:
            today = datetime.utcnow().date()
            week_start = (today - timedelta(days=today.weekday())).isoformat()
        data = report_service.get_weekly_summary(week_start)
        by_therapist = []
        for tname, tdata in data.get('by_therapist', {}).items():
            entry = {
                'therapist_id': tdata.get('therapist_id'),
                'therapist_name': tname,
                'patients': tdata.get('patients', []),
                'total_sessions': tdata.get('total_sessions', 0),
                'avg_score': tdata.get('avg_score', 0),
            }
            for p in entry['patients']:
                p['efficiency'] = round((p.get('avg_score', 0) or 0) * 0.5, 1)
            by_therapist.append(entry)
        return jsonify({
            'success': True,
            'data': {
                'week_start': data.get('week_start'),
                'week_end': data.get('week_end'),
                'by_therapist': by_therapist,
                'total_reports': data.get('total_reports', 0),
            }
        })
    except Exception as e:
        current_app.logger.error(f"Error in weekly-summary: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': 'Error al obtener resumen semanal'}), 500

@api_bp.route('/reports/accumulate', methods=['POST'])
@login_required
def api_reports_accumulate():
    if current_user.role != 'admin':
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403
    try:
        result = report_service.generate_all_weekly_reports()
        return jsonify({'success': True, 'message': f'Reportes acumulados: {len(result)} generados', 'count': len(result)})
    except Exception as e:
        current_app.logger.error(f"Error in reports/accumulate: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': 'Error al acumular reportes'}), 500

@api_bp.route('/reports/generate-all-weekly', methods=['POST'])
@login_required
def api_reports_generate_all_weekly():
    if current_user.role != 'admin':
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403
    try:
        week_start = request.args.get('week_start')
        result = report_service.generate_all_weekly_reports(week_start)
        return jsonify({'success': True, 'message': f'Se generaron {len(result)} reportes', 'count': len(result)})
    except Exception as e:
        current_app.logger.error(f"Error in reports/generate-all-weekly: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': 'Error al generar reportes'}), 500

@api_bp.route('/daily-reports', methods=['GET'])
@login_required
def api_daily_reports():
    if current_user.role != 'admin':
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403
    try:
        start = request.args.get('start')
        end = request.args.get('end')
        if not start or not end:
            today = datetime.utcnow().date()
            start = today.isoformat()
            end = today.isoformat()
        data = report_service.get_daily_reports(start, end)
        return jsonify({
            'success': True,
            'data': data,
        })
    except Exception as e:
        current_app.logger.error(f"Error in daily-reports: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': 'Error al obtener reportes diarios'}), 500

@api_bp.route('/therapist/efficiency', methods=['GET'])
@login_required
def api_therapist_efficiency():
    if current_user.role != 'admin':
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403
    try:
        therapist_id = request.args.get('therapist_id', type=int)
        from app.services.dashboard_service import DashboardService
        ds = DashboardService()
        from app.models import User
        therapists = [User.query.get(therapist_id)] if therapist_id else User.query.filter_by(role='terapista', is_active=True).all()
        breakdown = []
        for t in therapists:
            if not t:
                continue
            eff = ds.get_therapist_efficiency(t.id)
            breakdown.append({
                'therapist_id': t.id,
                'therapist_name': t.username,
                'audit_score': eff.get('avg_audit_score', 0),
                'feedback_score': eff.get('avg_feedback_score', 0),
                'efficiency': eff.get('efficiency', 0),
            })
        return jsonify({
            'success': True,
            'breakdown': breakdown,
        })
    except Exception as e:
        current_app.logger.error(f"Error in therapist/efficiency: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': 'Error al obtener eficiencia'}), 500

@api_bp.route('/sessions/current', methods=['GET'])
@login_required
def api_current_session():
    if current_user.role not in ('terapista', 'admin'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403
    now = datetime.now(timezone.utc).replace(tzinfo=None)
    appt = Appointment.query.filter(
        Appointment.therapist_id == current_user.id,
        Appointment.start_time <= now,
        Appointment.end_time >= now,
        Appointment.status.in_(['scheduled', 'in_progress'])
    ).order_by(Appointment.start_time).first()
    if not appt:
        return jsonify({'success': False, 'has_active': False})
    return jsonify({
        'success': True,
        'has_active': True,
        'session': {
            'id': appt.id,
            'title': appt.title or 'Sesión',
            'start': appt.start_time.isoformat() if appt.start_time else None,
            'end': appt.end_time.isoformat() if appt.end_time else None,
            'status': appt.status,
            'patient': {'id': appt.patient.id, 'name': appt.patient.username} if appt.patient else None,
            'location': appt.location,
        }
    })

@api_bp.route('/therapist/dashboard', methods=['GET'])
@login_required
def get_therapist_dashboard():
    if current_user.role != 'terapista':
        return jsonify({'error': 'Unauthorized'}), 403

    from datetime import datetime, timedelta
    from app.models import Appointment, User, SessionAudit
    from sqlalchemy import func as sqlfunc
    from app.extensions import db

    now = datetime.now()
    today = now.replace(hour=0, minute=0, second=0, microsecond=0)
    tomorrow = today + timedelta(days=1)
    
    # 1. Agenda de Hoy
    today_sessions = Appointment.query.filter(
        Appointment.therapist_id == current_user.id,
        Appointment.start_time >= today,
        Appointment.start_time < tomorrow,
        Appointment.status != 'cancelled'
    ).order_by(Appointment.start_time).all()

    agenda = []
    next_session = None
    
    for s in today_sessions:
        patient = User.query.get(s.patient_id) if s.patient_id else None
        is_current = s.start_time <= now and (s.end_time is None or s.end_time > now)
        session_info = {
            'id': s.id,
            'title': s.title or 'Sesión de Terapia',
            'patient': patient.username if patient else 'N/A',
            'start': s.start_time.strftime('%I:%M %p'),
            'location': s.location or '',
            'status': s.status,
            'is_current': is_current
        }
        agenda.append(session_info)
        if not next_session and (is_current or s.start_time > now):
            next_session = session_info

    # 2. Temas de la Sesión (Planned Text)
    planned_text = ""
    session_progress = 0
    avg_compliance = 0
    
    # Audit compliance data
    try:
        avg_cmp = db.session.query(sqlfunc.avg(SessionAudit.audit_score)).join(
            Appointment, SessionAudit.appointment_id == Appointment.id
        ).filter(
            Appointment.therapist_id == current_user.id,
            SessionAudit.audit_score.isnot(None)
        ).scalar()
        avg_compliance = float(avg_cmp) if avg_cmp else 0.0
    except Exception:
        pass

    if next_session:
        audit = SessionAudit.query.filter_by(appointment_id=next_session['id']).first()
        if audit and audit.planned_text:
            planned_text = audit.planned_text
            session_progress = int(audit.audit_score) if audit.audit_score else 0

    return jsonify({
        'success': True,
        'data': {
            'next_session': next_session,
            'agenda': agenda,
            'planned_text': planned_text,
            'session_progress': session_progress,
            'avg_compliance': avg_compliance,
            'total_students': len(set([s.patient_id for s in today_sessions]))
        }
    })
