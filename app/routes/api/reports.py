from app.routes.api._shared import (
    db, User, Notification, Appointment, Message, Game, SessionMetrics,
    SessionImage, ContactMessage, Sede, Payment, json, os, time, warnings,
    genai, Groq, _ollama_client, predict_level, start_async_training,
    get_user_today_utc_range, get_user_now, localize_datetime_for_display,
    get_user_timezone, bcrypt, limiter, csrf, EmailService, api_response,
    AvailabilityService, requests, or_, func,
    report_service,
    LIMA_TZ, _parse_json, _parse_datetime, analyze_contact_message_ai,
    AssignTherapistSchema, UpdateUserSchema, SendMessageSchema,
    uuid, secure_filename, datetime, timedelta, timezone,
    login_required, current_user, request, jsonify, current_app, url_for,
)
from app.routes.api import api_bp
@api_bp.route('/sessions/<int:appointment_id>/report-docx', methods=['GET'])
@login_required
def download_report_docx(appointment_id):
    if current_user.role not in ('terapista', 'admin', 'supervisor'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

    from app.models import SessionAudit, Appointment
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    appt = Appointment.query.get(appointment_id)
    if not appt:
        return jsonify({'success': False, 'error': 'Sesión no encontrada'}), 404

    if current_user.role == 'terapista' and appt.therapist_id != current_user.id:
        return jsonify({'success': False, 'error': 'No tienes acceso a esta sesión'}), 403

    audit = SessionAudit.query.filter_by(appointment_id=appointment_id).first()
    if not audit or audit.audit_status != 'completed':
        return jsonify({'success': False, 'error': 'No hay auditoría completada para esta sesión'}), 400

    report = audit.get_report()
    patient_name = appt.patient.username if appt.patient else '—'
    therapist_name = appt.therapist.username if appt.therapist else '—'

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    title = doc.add_heading('Reporte de Auditoría IA', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Sesión: {appt.title or "Sin título"}')
    doc.add_paragraph(f'Paciente: {patient_name}')
    doc.add_paragraph(f'Terapeuta: {therapist_name}')
    doc.add_paragraph(f'Fecha: {appt.start_time.strftime("%d/%m/%Y %H:%M") if appt.start_time else "—"}')
    doc.add_paragraph(f'Puntaje: {audit.audit_score or "—"} / 100')
    doc.add_paragraph(f'Clasificación: {report.get("status", "—")}')

    doc.add_heading('Objetivos Terapéuticos', level=1)
    objectives = report.get('objectives', [])
    if objectives:
        for obj in objectives:
            p = doc.add_paragraph()
            run = p.add_run(f'{obj.get("name", "—")}')
            run.bold = True
            p.add_run(f'\n  Evidencia: {obj.get("evidence", "—")}')
            p.add_run(f'\n  Clasificación: {obj.get("classification", "—")}')
    else:
        doc.add_paragraph('No se registraron objetivos.')

    if report.get('observations'):
        doc.add_heading('Observaciones IA', level=1)
        doc.add_paragraph(report['observations'])

    if audit.transcript_text:
        doc.add_heading('Transcripción', level=1)
        doc.add_paragraph(audit.transcript_text[:2000] + ('...' if len(audit.transcript_text) > 2000 else ''))

    feedback_items = []
    if audit.feedback_engagement:
        feedback_items.append(f'Compromiso: {audit.feedback_engagement}/5')
    if audit.feedback_progress:
        feedback_items.append(f'Progreso: {audit.feedback_progress}/5')
    if audit.feedback_notes:
        feedback_items.append(f'Notas: {audit.feedback_notes}')
    if feedback_items:
        doc.add_heading('Feedback del Terapeuta', level=1)
        for item in feedback_items:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph('Generado por EdySync', style='Intense Quote')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    patient_safe = patient_name.replace(' ', '_').lower()
    filename = f'auditoria_{patient_safe}_{appointment_id}.docx'

    from flask import send_file
    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )

@api_bp.route('/reports/generate-weekly', methods=['POST'])
@login_required
def generate_weekly_report():

    if current_user.role not in ('terapista', 'admin', 'supervisor'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

    data = request.get_json(silent=True) or {}
    patient_id = data.get('patient_id')
    week_start = data.get('week_start')

    if not patient_id or not week_start:
        return jsonify({'success': False, 'error': 'patient_id y week_start son requeridos'}), 400

    try:
        from app.services.report_service import ReportService
        rs = ReportService()
        therapist_id = current_user.id if current_user.role == 'terapista' else data.get('therapist_id', current_user.id)
        report = rs.generate_patient_weekly_report(patient_id, therapist_id, week_start)
        return jsonify({'success': True, 'report': report})
    except ValueError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        current_app.logger.error(f"Error generating weekly report: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@api_bp.route('/reports/weekly/<int:patient_id>', methods=['GET'])
@login_required
def get_weekly_report(patient_id):
    if current_user.role not in ('terapista', 'admin', 'supervisor', 'jugador'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

    week_start = request.args.get('week')

    try:
        from app.services.report_service import ReportService
        rs = ReportService()
        report = rs.get_patient_weekly_report(patient_id, week_start)
        if not report:
            return jsonify({'success': True, 'exists': False})
        return jsonify({'success': True, 'exists': True, 'report': report})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@api_bp.route('/reports/generate-daily', methods=['POST'])
@login_required
def generate_daily_report():

    if current_user.role not in ('terapista', 'admin', 'supervisor'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

    data = request.get_json(silent=True) or {}
    patient_id = data.get('patient_id')
    report_date = data.get('date')

    if not patient_id:
        return jsonify({'success': False, 'error': 'patient_id es requerido'}), 400

    try:
        from app.services.report_service import ReportService
        rs = ReportService()
        therapist_id = current_user.id if current_user.role == 'terapista' else data.get('therapist_id', current_user.id)
        report = rs.generate_daily_report(patient_id, therapist_id, report_date)
        return jsonify({'success': True, 'report': report})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@api_bp.route('/admin/audit-stats', methods=['GET'])
@login_required
def get_audit_stats():
    if current_user.role not in ('admin', 'supervisor'):
        return jsonify({'error': 'Unauthorized'}), 403
        
    try:
        from app.models import SessionAudit, Appointment, User
        from sqlalchemy import func
        from app.extensions import db

        total_audits = SessionAudit.query.filter(SessionAudit.audit_score.isnot(None)).count()
        avg_score_q = db.session.query(func.avg(SessionAudit.audit_score)).filter(SessionAudit.audit_score.isnot(None)).scalar()
        avg_score = float(avg_score_q) if avg_score_q else 0.0
        
        recent_audits = db.session.query(
            SessionAudit.id,
            SessionAudit.audit_score,
            SessionAudit.audited_at,
            Appointment.title,
            User.username.label('therapist_name')
        ).join(
            Appointment, SessionAudit.appointment_id == Appointment.id
        ).join(
            User, Appointment.therapist_id == User.id
        ).filter(
            SessionAudit.audit_score.isnot(None)
        ).order_by(SessionAudit.audited_at.desc()).limit(10).all()
        
        audit_rows = []
        for r in recent_audits:
            audit_rows.append({
                'id': r[0],
                'score': float(r[1]),
                'date': r[2].isoformat() if r[2] else None,
                'title': r[3],
                'therapist': r[4]
            })

        therapist_scores = db.session.query(
            User.username,
            func.avg(SessionAudit.audit_score),
            func.count(SessionAudit.id)
        ).join(
            Appointment, SessionAudit.appointment_id == Appointment.id
        ).join(
            User, Appointment.therapist_id == User.id
        ).filter(
            SessionAudit.audit_score.isnot(None)
        ).group_by(User.id).all()
        
        therapist_stats = []
        for t in therapist_scores:
            therapist_stats.append({
                'name': t[0],
                'avg_score': round(float(t[1]), 1),
                'count': t[2]
            })

        return jsonify({
            'success': True,
            'data': {
                'total': total_audits,
                'avg_score': round(avg_score, 1),
                'recent': audit_rows,
                'by_therapist': therapist_stats
            }
        })
    except Exception as e:
        current_app.logger.warning(f"Failed to load audit stats: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@api_bp.route('/therapist/weekly-reports/pending', methods=['GET'])
@login_required
def api_weekly_reports_pending():
    try:
        therapist_id = current_user.id
        if current_user.role == 'terapista':
            pass
        elif current_user.role in ('admin', 'supervisor'):
            therapist_id = request.args.get('therapist_id', type=int) or therapist_id
        else:
            return jsonify({'success': False, 'error': 'Acceso denegado'}), 403

        from app.models import WeeklyReport, Notification
        from sqlalchemy import inspect as sa_inspect
        from app.extensions import db
        inspector = sa_inspect(db.engine)
        if 'weekly_report' not in inspector.get_table_names():
            db.create_all()
            current_app.logger.info("Auto-migration: created tables on-demand")

        week_start = datetime.utcnow().date()
        monday = week_start - timedelta(days=week_start.weekday())
        reports = WeeklyReport.query.filter(
            WeeklyReport.therapist_id == therapist_id,
            WeeklyReport.week_start == monday
        ).count()
        notification = Notification.query.filter(
            Notification.user_id == current_user.id,
            Notification.type == 'reportes',
            Notification.is_read == False
        ).order_by(Notification.timestamp.desc()).first()
        return jsonify({
            'success': True,
            'has_pending': reports > 0,
            'reports_count': reports,
            'has_notification': notification is not None,
            'week_start': monday.isoformat(),
            'week_end': (monday + timedelta(days=6)).isoformat(),
        })
    except Exception as e:
        current_app.logger.error(f"Error in weekly-reports/pending: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': 'Error al consultar reportes'}), 500

@api_bp.route('/therapist/weekly-reports/generate', methods=['POST'])
@login_required
def api_weekly_reports_generate():
    if current_user.role not in ('terapista', 'admin', 'supervisor'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403
    try:
        therapist_id = current_user.id
        today = datetime.utcnow().date()
        monday = today - timedelta(days=today.weekday())
        week_start = monday.isoformat()

        patients = current_user.associated_patients.filter_by(role='jugador', is_active=True).all() if hasattr(current_user, 'associated_patients') else []

        from app.services.report_service import ReportService
        rs = ReportService()
        generated = []
        for patient in patients:
            try:
                report = rs.generate_patient_weekly_report(patient.id, therapist_id, week_start)
                generated.append(report)
            except Exception as e:
                current_app.logger.warning(f"Weekly report error for patient {patient.id}: {e}")

        return jsonify({
            'success': True,
            'reports_count': len(generated),
            'patients_count': len(patients),
            'report': generated[0] if generated else None
        })
    except Exception as e:
        current_app.logger.error(f"Error in weekly-reports/generate: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': 'Error al generar reportes'}), 500

@api_bp.route('/weekly-summary', methods=['GET'])
@login_required
def api_weekly_summary():
    if current_user.role not in ('admin', 'supervisor'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403
    try:
        week_start = request.args.get('week_start')
        if not week_start:
            today = datetime.utcnow().date()
            week_start = (today - timedelta(days=today.weekday())).isoformat()
        data = report_service.get_weekly_summary(week_start)
        by_therapist = []
        for tname, tdata in data.get('by_therapist', {}).items():
            entry = {
                'therapist_id': tdata.get('therapist_id'),
                'therapist_name': tname,
                'patients': tdata.get('patients', []),
                'total_sessions': tdata.get('total_sessions', 0),
                'avg_score': tdata.get('avg_score', 0),
            }
            for p in entry['patients']:
                p['efficiency'] = round((p.get('avg_score', 0) or 0) * 0.5, 1)
            by_therapist.append(entry)
        return jsonify({
            'success': True,
            'data': {
                'week_start': data.get('week_start'),
                'week_end': data.get('week_end'),
                'by_therapist': by_therapist,
                'total_reports': data.get('total_reports', 0),
            }
        })
    except Exception as e:
        current_app.logger.error(f"Error in weekly-summary: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': 'Error al obtener resumen semanal'}), 500

@api_bp.route('/reports/accumulate', methods=['POST'])
@login_required
def api_reports_accumulate():
    if current_user.role not in ('admin', 'supervisor'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403
    try:
        result = report_service.generate_all_weekly_reports()
        return jsonify({'success': True, 'message': f'Reportes acumulados: {len(result)} generados', 'count': len(result)})
    except Exception as e:
        current_app.logger.error(f"Error in reports/accumulate: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': 'Error al acumular reportes'}), 500

@api_bp.route('/reports/generate-all-weekly', methods=['POST'])
@login_required
def api_reports_generate_all_weekly():
    if current_user.role not in ('admin', 'supervisor'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403
    try:
        week_start = request.args.get('week_start')
        result = report_service.generate_all_weekly_reports(week_start)
        return jsonify({'success': True, 'message': f'Se generaron {len(result)} reportes', 'count': len(result)})
    except Exception as e:
        current_app.logger.error(f"Error in reports/generate-all-weekly: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': 'Error al generar reportes'}), 500

@api_bp.route('/daily-reports', methods=['GET'])
@login_required
def api_daily_reports():
    if current_user.role not in ('admin', 'supervisor'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403
    try:
        start = request.args.get('start')
        end = request.args.get('end')
        if not start or not end:
            today = datetime.utcnow().date()
            start = today.isoformat()
            end = today.isoformat()
        data = report_service.get_daily_reports(start, end)
        return jsonify({
            'success': True,
            'data': data,
        })
    except Exception as e:
        current_app.logger.error(f"Error in daily-reports: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': 'Error al obtener reportes diarios'}), 500

@api_bp.route('/therapist/efficiency', methods=['GET'])
@login_required
def api_therapist_efficiency():
    if current_user.role not in ('admin', 'supervisor'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403
    try:
        therapist_id = request.args.get('therapist_id', type=int)
        from app.services.dashboard_service import DashboardService
        ds = DashboardService()
        from app.models import User
        therapists = [User.query.get(therapist_id)] if therapist_id else User.query.filter_by(role='terapista', is_active=True).all()
        breakdown = []
        for t in therapists:
            if not t:
                continue
            eff = ds.get_therapist_efficiency(t.id)
            breakdown.append({
                'therapist_id': t.id,
                'therapist_name': t.username,
                'audit_score': eff.get('avg_audit_score', 0),
                'feedback_score': eff.get('avg_feedback_score', 0),
                'efficiency': eff.get('efficiency', 0),
            })
        return jsonify({
            'success': True,
            'breakdown': breakdown,
        })
    except Exception as e:
        current_app.logger.error(f"Error in therapist/efficiency: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': 'Error al obtener eficiencia'}), 500

@api_bp.route('/reports/monthly', methods=['GET'])
@login_required
def api_reports_monthly():
    if current_user.role not in ('admin', 'supervisor'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403
    try:
        year = request.args.get('year', type=int, default=datetime.utcnow().year)
        month = request.args.get('month', type=int, default=datetime.utcnow().month)
        summary = report_service.get_monthly_summary(year, month)
        return jsonify({'success': True, 'summary': summary})
    except Exception as e:
        current_app.logger.error(f"Error in reports/monthly: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': 'Error al obtener reporte mensual'}), 500

@api_bp.route('/reports/quarterly', methods=['GET'])
@login_required
def api_reports_quarterly():
    if current_user.role not in ('admin', 'supervisor'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403
    try:
        year = request.args.get('year', type=int, default=datetime.utcnow().year)
        quarter = request.args.get('quarter', type=int, default=(datetime.utcnow().month - 1) // 3 + 1)
        summary = report_service.get_quarterly_summary(year, quarter)
        return jsonify({'success': True, 'summary': summary})
    except Exception as e:
        current_app.logger.error(f"Error in reports/quarterly: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': 'Error al obtener reporte trimestral'}), 500

@api_bp.route('/reports/generate-monthly', methods=['POST'])
@login_required
@csrf.exempt
def api_reports_generate_monthly():
    if current_user.role not in ('admin', 'supervisor'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403
    try:
        year = request.args.get('year', type=int, default=datetime.utcnow().year)
        month = request.args.get('month', type=int, default=datetime.utcnow().month)
        generated = report_service.generate_all_monthly_reports(year, month)
        return jsonify({
            'success': True,
            'message': f'{len(generated)} reportes mensuales generados',
            'count': len(generated)
        })
    except Exception as e:
        current_app.logger.error(f"Error in reports/generate-monthly: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': 'Error al generar reportes mensuales'}), 500

@api_bp.route('/reports/generate-quarterly', methods=['POST'])
@login_required
@csrf.exempt
def api_reports_generate_quarterly():
    if current_user.role not in ('admin', 'supervisor'):
        return jsonify({'success': False, 'error': 'Acceso denegado'}), 403
    try:
        year = request.args.get('year', type=int, default=datetime.utcnow().year)
        quarter = request.args.get('quarter', type=int, default=(datetime.utcnow().month - 1) // 3 + 1)
        generated = report_service.generate_all_quarterly_reports(year, quarter)
        return jsonify({
            'success': True,
            'message': f'{len(generated)} reportes trimestrales generados',
            'count': len(generated)
        })
    except Exception as e:
        current_app.logger.error(f"Error in reports/generate-quarterly: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': 'Error al generar reportes trimestrales'}), 500

