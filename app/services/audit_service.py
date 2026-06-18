import json
import logging
import os
from datetime import datetime

logger = logging.getLogger(__name__)

MAX_AUDIT_PROMPT_CHARS = 1800


def _truncate_audit_text(text, max_chars=MAX_AUDIT_PROMPT_CHARS):
    """Recorta texto para caber en el límite de tokens de Groq (inicio + final)."""
    if not text or len(text) <= max_chars:
        return text
    half = max_chars // 2
    return text[:half] + '\n\n[... contenido omitido por límite de análisis ...]\n\n' + text[-half:]


_STOP_WORDS_ES = {
    'de',
    'la',
    'que',
    'el',
    'en',
    'y',
    'a',
    'los',
    'del',
    'se',
    'las',
    'por',
    'un',
    'para',
    'con',
    'no',
    'una',
    'su',
    'al',
    'lo',
    'como',
    'más',
    'pero',
    'sus',
    'le',
    'ya',
    'este',
    'entre',
    'porque',
    'esta',
    'muy',
    'todo',
    'sin',
    'ello',
    'cada',
    'otro',
    'cual',
    'cuando',
    'donde',
    'quien',
    'aquel',
    'solo',
    'allí',
    'así',
    'tras',
    'entonces',
    'tiempo',
    'también',
    'sea',
    'sido',
    'han',
    'ser',
    'haber',
    'tener',
    'hacer',
    'estar',
    'poder',
    'ir',
    'dar',
    'ver',
    'decir',
    'saber',
    'creer',
    'era',
    'son',
    'fue',
    'has',
    'había',
    'hay',
}


def extract_docx_text(file_path):
    """Extrae texto de .docx usando python-docx"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError('python-docx no está instalado. Ejecuta: pip install python-docx')

    if not os.path.exists(file_path):
        raise ValueError(f'Archivo no encontrado: {file_path}')

    try:
        doc = Document(file_path)
    except Exception as e:
        raise ValueError(f'Error al abrir el archivo Word: {str(e)}')

    sections = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style and para.style.name.startswith('Heading'):
            sections.append(f'\n## {text}')
        else:
            sections.append(text)

    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(' | '.join(cells))
        if table_rows:
            sections.append('\n[TABLA]\n' + '\n'.join(table_rows) + '\n[/TABLA]')

    full_text = '\n'.join(sections)

    if not full_text.strip():
        raise ValueError('El documento Word está vacío o no contiene texto extraíble')

    logger.info(f'Texto extraído del Word: {len(full_text)} caracteres')
    return full_text


def transcribe_audio(file_path):
    """Transcribe audio con Whisper (se borra tras transcripción)"""
    api_key = os.getenv('GROQ_API_KEY')
    if not api_key:
        raise ValueError('GROQ_API_KEY no está configurada en las variables de entorno')

    if not os.path.exists(file_path):
        raise ValueError(f'Archivo de audio no encontrado: {file_path}')

    try:
        from groq import Groq

        client = Groq(api_key=api_key)

        with open(file_path, 'rb') as audio_file:
            transcription = client.audio.transcriptions.create(
                file=(os.path.basename(file_path), audio_file),
                model='whisper-large-v3-turbo',
                language='es',
                response_format='verbose_json',
                temperature=0.0,
            )

        transcript_text = transcription.text or ''
        duration = getattr(transcription, 'duration', None)
        language = getattr(transcription, 'language', 'es')

        logger.info(f'Audio transcrito: {len(transcript_text)} caracteres, duración={duration}s')

        return {'text': transcript_text, 'duration': int(duration) if duration else 0, 'language': language}

    except Exception as e:
        logger.error(f'Error en transcripción Whisper/Groq: {str(e)}')
        raise ValueError(f'Error al transcribir el audio: {str(e)}')

    finally:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f'Audio eliminado por privacidad: {file_path}')
        except Exception as del_err:
            logger.error(f'ERROR CRITICO: No se pudo eliminar el audio {file_path}: {del_err}')


AUDIT_SYSTEM_PROMPT = """Eres un auditor clínico experto del Centro de Terapias Juan Pablo II.
Tu tarea es comparar la PROGRAMACIÓN PLANIFICADA (objetivos terapéuticos del documento Word)
con la TRANSCRIPCIÓN REAL de la sesión (audio transcrito por Whisper).

INSTRUCCIONES ESTRICTAS:
1. Identifica TODOS los objetivos terapéuticos del documento de programación.
2. Busca evidencia en la transcripción de que cada objetivo fue abordado.
3. Clasifica CADA objetivo como:
   - "logrado": Se abordó completamente según lo planificado
   - "parcial": Se mencionó o intentó pero no se completó
   - "no_cubierto": No hay evidencia de que se haya trabajado
4. Calcula un score de cumplimiento (0-100).
5. Identifica actividades extra que se realizaron pero no estaban planificadas.

SIEMPRE responde en este formato JSON exacto (sin texto adicional):
{
  "score": 85.0,
  "status": "cumple_parcial",
  "objectives": [
    {
      "name": "Nombre del objetivo",
      "classification": "logrado|parcial|no_cubierto",
      "evidence": "Extracto relevante de la transcripción que lo respalda"
    }
  ],
  "planned_activities": ["actividad 1", "actividad 2"],
  "executed_activities": ["actividad ejecutada 1"],
  "missing_activities": ["actividad no ejecutada"],
  "extra_activities": ["actividad no planificada que sí se realizó"],
  "observations": "Observaciones generales del auditor",
  "recommendations": ["Recomendación 1", "Recomendación 2"]
}

REGLAS para el campo "status":
- "cumple": score >= 80
- "cumple_parcial": score >= 50 y < 80
- "no_cumple": score < 50
"""


def run_audit(appointment_id):
    """Auditoría IA completa: compara planned_text vs transcript_text con Groq"""
    from app.models import SessionAudit, SessionImage, db

    audit = SessionAudit.query.filter_by(appointment_id=appointment_id).first()
    if not audit:
        raise ValueError('No existe registro de auditoría para esta sesión')

    if not audit.planned_text:
        raise ValueError('No se ha subido la programación (.docx) para esta sesión')

    if not audit.transcript_text:
        raise ValueError('No se ha transcrito el audio de esta sesión')

    audit.audit_status = 'processing'
    db.session.commit()

    api_key = os.getenv('GROQ_API_KEY')
    if not api_key:
        audit.audit_status = 'error'
        db.session.commit()
        raise ValueError('GROQ_API_KEY no configurada')

    images = SessionImage.query.filter_by(appointment_id=appointment_id).all()
    photos_context = ''
    if images:
        photo_descriptions = [
            f"- Foto {i + 1}: tipo={img.image_type}, notas='{img.notes or 'sin notas'}'" for i, img in enumerate(images)
        ]
        photos_context = f'\n\nFOTOS DE LA SESIÓN ({len(images)} archivos):\n' + '\n'.join(photo_descriptions)

    planned_for_prompt = _truncate_audit_text(audit.planned_text)
    transcript_for_prompt = _truncate_audit_text(audit.transcript_text)

    user_prompt = f"""PROGRAMACIÓN PLANIFICADA (extraída del documento Word):
---
{planned_for_prompt}
---

TRANSCRIPCIÓN REAL DE LA SESIÓN (audio transcrito por Whisper):
---
{transcript_for_prompt}
---
{photos_context}

Analiza y genera el reporte de cumplimiento en formato JSON."""

    try:
        from groq import Groq

        client = Groq(api_key=api_key)

        response = client.chat.completions.create(
            model='llama-3.1-8b-instant',
            messages=[{'role': 'system', 'content': AUDIT_SYSTEM_PROMPT}, {'role': 'user', 'content': user_prompt}],
            temperature=0.1,
            max_tokens=2000,
            response_format={'type': 'json_object'},
        )

        raw_content = (response.choices[0].message.content or '').strip()

        report = json.loads(raw_content)

        audit.audit_report_json = json.dumps(report, ensure_ascii=False)
        audit.audit_score = float(report.get('score', 0))
        audit.audit_status = 'completed'
        audit.audited_at = datetime.utcnow()
        db.session.commit()

        objectives = report.get('objectives', [])
        if objectives:
            weight_map = {'logrado': 1.0, 'parcial': 0.5, 'no_cubierto': 0.0}
            scores = [weight_map.get(o.get('classification', 'no_cubierto'), 0.0) * 100 for o in objectives]
            real_score = sum(scores) / len(scores)
            real_score = round(real_score, 1)
            if abs(real_score - audit.audit_score) > 10:
                logger.info(f'Score recalculado por objetivos: {audit.audit_score} → {real_score}')
                audit.audit_score = real_score
                report['score'] = real_score
                if real_score >= 80:
                    report['status'] = 'cumple'
                elif real_score >= 50:
                    report['status'] = 'cumple_parcial'
                else:
                    report['status'] = 'no_cumple'
                audit.audit_report_json = json.dumps(report, ensure_ascii=False)
                db.session.commit()

        vectorial = compute_similarity_vectorial(audit.planned_text, audit.transcript_text)
        score_vectorial = vectorial['score_vectorial']

        duracion = audit.audio_duration_seconds or 0
        DURACION_ESPERADA = 2700
        ratio = min(1.0, duracion / DURACION_ESPERADA)
        factor = min(1.0, ratio / 0.1)

        score_final = round((audit.audit_score * 0.5 + score_vectorial * 0.5) * factor, 1)
        logger.info(
            f'Score final: LLM={audit.audit_score} vectorial={score_vectorial} factor={factor:.2f} → {score_final}'
        )
        audit.audit_score = score_final
        report['score'] = score_final
        if score_final >= 80:
            report['status'] = 'cumple'
        elif score_final >= 50:
            report['status'] = 'cumple_parcial'
        else:
            report['status'] = 'no_cumple'
        audit.audit_report_json = json.dumps(report, ensure_ascii=False)
        db.session.commit()

        logger.info(f'Auditoría completada para sesión {appointment_id}: score={audit.audit_score}')
        return report

    except json.JSONDecodeError as e:
        audit.audit_status = 'error'
        raw_preview = (locals().get('raw_content') or '')[:500]
        audit.audit_report_json = json.dumps({'error': 'El LLM no devolvió JSON válido', 'raw_response': raw_preview})
        db.session.commit()
        logger.error(f'Error parseando JSON de Llama: {e}')
        raise ValueError(f'Error al procesar respuesta de IA: {str(e)}')

    except Exception as e:
        audit.audit_status = 'error'
        db.session.commit()
        logger.error(f'Error en auditoría IA: {str(e)}')
        raise ValueError(f'Error en auditoría IA: {str(e)}')


ATTENDANCE_SYSTEM_PROMPT = """Eres un asistente que determina si un paciente asistió a su sesión terapéutica.
Comparas el PLAN DE SESIÓN con la TRANSCRIPCIÓN REAL.

Reglas:
1. Si la transcripción tiene menos de 50 caracteres o solo contiene ruido/saludos sin contenido → el paciente NO asistió.
2. Si la transcripción cubre al menos un 5% de las actividades planificadas → el paciente SÍ asistió.
3. Calcula el porcentaje de cobertura: qué tanto de lo planificado se menciona/ejecuta en la transcripción.

Responde SOLO con JSON:
{
  "suggested_attendance": "present" o "absent",
  "confidence": 0.0-1.0,
  "coverage_pct": 0-100,
  "reason": "Explicación breve"
}
"""


def analyze_attendance(planned_text, transcript_text):
    """Analiza si el paciente asistió comparando plan vs transcripción"""
    api_key = os.getenv('GROQ_API_KEY')
    if not api_key:
        raise ValueError('GROQ_API_KEY no está configurada')

    try:
        from groq import Groq

        client = Groq(api_key=api_key)

        user_prompt = f"""PLAN DE SESIÓN:
---
{_truncate_audit_text(planned_text)}
---

TRANSCRIPCIÓN REAL:
---
{_truncate_audit_text(transcript_text)}
---

Analiza si el paciente asistió y cubrió al menos el 5% de lo planificado."""

        response = client.chat.completions.create(
            model='llama-3.1-8b-instant',
            messages=[
                {'role': 'system', 'content': ATTENDANCE_SYSTEM_PROMPT},
                {'role': 'user', 'content': user_prompt},
            ],
            temperature=0.1,
            max_tokens=500,
            response_format={'type': 'json_object'},
        )

        import json

        result = json.loads(response.choices[0].message.content.strip())

        return {
            'suggested_attendance': result.get('suggested_attendance', 'present'),
            'confidence': float(result.get('confidence', 0.5)),
            'coverage_pct': float(result.get('coverage_pct', 50)),
            'reason': result.get('reason', ''),
        }

    except Exception as e:
        logger.error(f'Error in analyze_attendance: {str(e)}')
        if len(transcript_text.strip()) < 50:
            return {
                'suggested_attendance': 'absent',
                'confidence': 0.8,
                'coverage_pct': 0,
                'reason': 'Transcripción insuficiente',
            }
        return {
            'suggested_attendance': 'present',
            'confidence': 0.6,
            'coverage_pct': 50,
            'reason': 'Fallback: no se pudo analizar con IA',
        }


def _dividir_objetivos(texto):
    lineas = texto.strip().split('\n')
    objetivos = []
    for l in lineas:
        l = l.strip()
        if len(l) > 15 and not l.startswith('#'):
            objetivos.append(l)
    if not objetivos:
        objetivos = [texto[:500]]
    return objetivos[:20]


def compute_similarity_vectorial(planned_text, transcript_text):
    import numpy as np
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity

    if not transcript_text or not planned_text:
        return {'score_vectorial': 0, 'objetivos_cubiertos': 0, 'n_objectives': 0}

    objetivos = _dividir_objetivos(planned_text)
    if not objetivos:
        return {'score_vectorial': 0, 'objetivos_cubiertos': 0, 'n_objectives': 0}

    docs = objetivos + [transcript_text]
    try:
        vectorizer = TfidfVectorizer(max_features=1000, stop_words=list(_STOP_WORDS_ES))
        tfidf = vectorizer.fit_transform(docs)
        objetivo_vecs = tfidf[:-1]
        transcript_vec = tfidf[-1:]

        sims = cosine_similarity(objetivo_vecs, transcript_vec).flatten()
        score = float(np.mean(sims) * 100)
        cubiertos = int(np.sum(sims > 0.15))
        return {
            'score_vectorial': round(min(score, 100), 1),
            'objetivos_cubiertos': cubiertos,
            'n_objectives': len(objetivos),
        }
    except Exception as e:
        logger.error(f'Error en vectorización: {e}')
        return {'score_vectorial': 50, 'objetivos_cubiertos': 0, 'n_objectives': len(objetivos)}
